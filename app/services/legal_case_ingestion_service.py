"""
Legal Case Ingestion Service

Complete pipeline for ingesting historical legal cases into the knowledge management system.
Handles file upload, text extraction, section segmentation, and database storage.
"""

import os
import re
import hashlib
import logging
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime, date
from pathlib import Path

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

# PDF/DOCX extraction
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Enhanced PDF processor for better Arabic text extraction
try:
    from ..processors.enhanced_arabic_pdf_processor import EnhancedArabicPDFProcessor
    ENHANCED_PDF_AVAILABLE = True
except ImportError:
    ENHANCED_PDF_AVAILABLE = False
    logger.warning("EnhancedArabicPDFProcessor not available, using basic extraction")

from ..models.legal_knowledge import (
    KnowledgeDocument, LegalCase, CaseSection, KnowledgeChunk
)

logger = logging.getLogger(__name__)


class LegalCaseIngestionService:
    """Service for ingesting legal cases from PDF/DOCX/TXT files."""
    
    def __init__(self, db: AsyncSession, upload_dir: str = "uploads/legal_cases"):
        """
        Initialize the legal case ingestion service.
        
        Args:
            db: Async database session
            upload_dir: Directory to store uploaded files
        """
        self.db = db
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Enhanced Arabic section keywords for segmentation
        # Patterns are ordered from most specific to least specific for better matching
        self.section_patterns = {
            'summary': [
                # Introductory Islamic phrases (common in formal documents)
                r'الحمدلله\s+والصلاة\s+والسلام\s+على\s+رسول\s+الله',
                r'الحمد\s+لله\s+والصلاة\s+والسلام',
                r'أما\s+بعد',
                # Standard summary keywords
                r'ملخص\s+القضية',
                r'ملخص\s+الدعوى',
                r'ملخص',
                r'نبذة',
                r'موجز',
                r'الملخص'
            ],
            'facts': [
                # Specific fact introduction phrases
                r'تتحصل\s+وقائع',
                r'ما\s+ورد\s+في\s+صحيفة\s+الدعوى',
                r'ما\s+جاء\s+في\s+الدعوى',
                r'وقائع\s+القضية',
                r'وقائع\s+الدعوى',
                # General fact keywords
                r'الوقائع',
                r'الواقعة',
                r'الحادثة'
            ],
            'arguments': [
                # Legal arguments and pleadings
                r'تقرير\s+المحامي',
                r'طلب\s+إلزام',
                r'حجج\s+الأطراف',
                r'أقوال\s+الأطراف',
                r'المرافعات',
                r'الدفوع',
                # General argument keywords
                r'الأسباب',
                r'الحجج',
                r'دفاع',
                r'الحجة'
            ],
            'ruling': [
                # Specific ruling phrases
                r'نص\s+الحكم',
                r'منطوق\s+الحكم',
                r'حكمت\s+المحكمة',
                r'قررت\s+المحكمة',
                r'حكم\s+نهائي',
                # General ruling keywords
                r'المنطوق',
                r'الحكم',
                r'القرار'
            ],
            'legal_basis': [
                # Specific legal references and articles
                r'نظام\s+المحاكم\s+التجارية',
                r'لائحة\s+الدعوى',
                r'أوراق\s+ومستندات',
                r'المادة\s+الثلاثون',
                r'المادة\s+الثلاثين',
                r'المادة\s+التاسعة\s+والعشرون',
                r'المادة\s+التاسعة\s+والعشرين',
                r'المادة\s+\d+',  # Matches "المادة" followed by any number
                # Legal basis phrases
                r'الأساس\s+القانوني',
                r'السند\s+القانوني',
                r'التكييف\s+القانوني',
                r'الأسانيد\s+القانونية',
                r'المستند\s+القانوني',
                r'الحيثيات'
            ]
        }
    
    # =====================================================
    # FILE UPLOAD AND HASH CALCULATION
    # =====================================================
    
    async def save_uploaded_case_file(
        self,
        file_content: bytes,
        filename: str,
        uploaded_by: int
    ) -> Tuple[str, str, KnowledgeDocument]:
        """
        Save uploaded file and create KnowledgeDocument record.
        
        Args:
            file_content: Binary content of the uploaded file
            filename: Original filename
            uploaded_by: User ID who uploaded the file
            
        Returns:
            Tuple of (file_path, file_hash, knowledge_document)
            
        Raises:
            ValueError: If file already exists (duplicate hash)
        """
        # Calculate SHA-256 hash
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        # Check for duplicates
        duplicate_check = await self.db.execute(
            select(KnowledgeDocument).where(
                KnowledgeDocument.file_hash == file_hash
            )
        )
        existing_doc = duplicate_check.scalar_one_or_none()
        
        if existing_doc:
            raise ValueError(
                f"Duplicate file detected. Document already exists: {existing_doc.title} (ID: {existing_doc.id})"
            )
        
        # Generate unique filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        file_extension = Path(filename).suffix
        safe_filename = f"{timestamp}_{file_hash[:12]}{file_extension}"
        file_path = self.upload_dir / safe_filename
        
        # Save file to disk
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        logger.info(f"Saved file to: {file_path}")
        
        # Create KnowledgeDocument record
        knowledge_doc = KnowledgeDocument(
            title=Path(filename).stem,  # Use filename without extension as title
            category='case',
            file_path=str(file_path),
            file_hash=file_hash,
            source_type='uploaded',
            status='raw',
            uploaded_by=uploaded_by,
            uploaded_at=datetime.utcnow(),
            document_metadata={
                'original_filename': filename,
                'file_size': len(file_content),
                'uploaded_by': uploaded_by,
                'file_type': Path(filename).suffix.lower()
            }
        )
        
        self.db.add(knowledge_doc)
        await self.db.flush()  # Get the ID
        
        logger.info(f"Created KnowledgeDocument ID: {knowledge_doc.id}")
        
        return str(file_path), file_hash, knowledge_doc
    
    # =====================================================
    # TEXT EXTRACTION
    # =====================================================
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from PDF, DOCX, or TXT file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text as a single string
            
        Raises:
            ValueError: If file format is not supported
            RuntimeError: If extraction fails
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self._extract_pdf_text(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._extract_docx_text(file_path)
        elif file_extension == '.txt':
            return self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """
        Extract text from PDF using advanced dict extraction (from extract_arabic_pdf.py).
        This method properly handles Arabic text with full fixing and RTL processing.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text with proper Arabic text direction
        """
        try:
            if not fitz:
                raise RuntimeError("PyMuPDF (fitz) not available")
            
            doc = fitz.open(str(file_path))
            text = ""
            
            # Save page count before processing
            total_pages = len(doc)
            logger.info(f"Starting advanced PDF extraction with dict method for {total_pages} pages")
            
            for page_num, page in enumerate(doc, 1):
                try:
                    logger.info(f"Processing page {page_num}/{total_pages} with dict extraction...")
                    
                    # Use get_text("dict") to get full structured text
                    page_dict = page.get_text("dict")
                    
                    if not page_dict or "blocks" not in page_dict:
                        logger.warning(f"No dict blocks found at page {page_num}")
                        text += f"\n---EMPTY_PAGE_{page_num}---\n"
                        continue
                    
                    blocks = page_dict["blocks"]
                    logger.info(f"Page {page_num}: Found {len(blocks)} blocks in dict")
                    
                    # Process each block: blocks -> lines -> spans
                    for block_num, block in enumerate(blocks):
                        if "lines" not in block:
                            text += "\n"
                            continue
                        
                        lines = block["lines"]
                        
                        # Process each line
                        for line_num, line in enumerate(lines):
                            if "spans" not in line:
                                text += "\n"
                                continue
                            
                            spans = line["spans"]
                            line_text = ""
                            
                            # Extract each span
                            for span in spans:
                                if "text" not in span:
                                    continue
                                
                                span_text = span["text"]
                                if span_text.strip():
                                    line_text += span_text
                            
                            # Apply advanced Arabic fixing on each line
                            if line_text.strip():
                                # Always apply fixing for Arabic text from PDFs
                                if self._needs_fixing(line_text):
                                    # This does: clean artifacts + normalize + reshape + BiDi
                                    fixed_line = self._fix_arabic_text(line_text)
                                    text += fixed_line + "\n"
                                else:
                                    # For non-Arabic or already good text
                                    text += line_text + "\n"
                            else:
                                text += "\n"
                    
                    # Add page separator
                    text += "\n---PAGE_SEPARATOR---\n"
                    
                except Exception as page_e:
                    logger.error(f"Error processing page {page_num}: {page_e}")
                    text += f"\n---ERROR_PAGE_{page_num}---\n"
                    continue
            
            # Close document before logging (to avoid accessing closed doc)
            doc.close()
            logger.info(f"✅ Extracted {len(text)} characters from {total_pages} pages using advanced dict extraction")
            return text
            
        except Exception as e:
            logger.error(f"Advanced PDF extraction failed: {e}")
            raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        if not DocxDocument:
            raise RuntimeError(
                "python-docx not installed. "
                "Install with: pip install python-docx"
            )
        
        try:
            doc = DocxDocument(str(file_path))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text
        """
        try:
            # Try UTF-8 first (most common for Arabic text)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                    logger.info(f"Extracted {len(text)} characters from TXT (UTF-8)")
                    return text
            except UnicodeDecodeError:
                # Fallback to other encodings for Arabic text
                logger.warning("UTF-8 decoding failed, trying Windows-1256 (Arabic)")
                try:
                    with open(file_path, 'r', encoding='windows-1256') as f:
                        text = f.read()
                        logger.info(f"Extracted {len(text)} characters from TXT (Windows-1256)")
                        return text
                except UnicodeDecodeError:
                    # Final fallback to ISO-8859-1 (latin-1)
                    logger.warning("Windows-1256 failed, trying ISO-8859-1")
                    with open(file_path, 'r', encoding='iso-8859-1') as f:
                        text = f.read()
                        logger.info(f"Extracted {len(text)} characters from TXT (ISO-8859-1)")
                        return text
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from TXT: {str(e)}")
    
    # =====================================================
    # ADVANCED ARABIC TEXT PROCESSING (from extract_arabic_pdf.py)
    # =====================================================
    
    def _needs_fixing(self, text: str) -> bool:
        """Enhanced detection of Arabic text that needs fixing - ALWAYS fix Arabic text in PDFs"""
        if not text.strip():
            return False

        # Check for any Arabic characters
        arabic_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF')
        if arabic_chars == 0:
            return False

        # Check for fragmented words (average word length <= 2)
        words = text.split()
        if words:
            avg_word_len = sum(len(w) for w in words) / len(words)
            if avg_word_len <= 2:
                return True

        # Check for fragmented individual Arabic characters
        fragmented_patterns = [
            len(word) == 1 and '\u0600' <= word <= '\u06FF' for word in words
        ]
        if any(fragmented_patterns):
            return True

        # Check for artifacts (isolated Unicode characters)
        artifacts = ['ﻢ', 'ﻪ', 'ﻆ', 'ﺍ', 'ﺕ', 'ﺏ', 'ﻞ', 'ﺝ', 'ﺡ', 'ﺥ', 'ﺩ', 'ﺫ', 'ﺭ', 'ﺯ', 'ﺱ', 'ﺵ', 'ﺹ', 'ﺽ', 'ﻁ', 'ﻅ', 'ﻉ', 'ﻍ', 'ﻑ', 'ﻕ', 'ﻙ', 'ﻝ', 'ﻡ', 'ﻥ', 'ﻩ', 'ﻭ', 'ﻱ']
        if any(artifact in text for artifact in artifacts):
            return True

        # For PDFs, ALWAYS apply fixing if Arabic text is detected
        arabic_ratio = arabic_chars / len(text.strip()) if text.strip() else 0
        if arabic_ratio > 0.1:  # If 10% or more Arabic characters, fix it
            return True

        return False

    def _clean_text_artifacts(self, text: str) -> str:
        """Remove artifacts and clean up text formatting"""
        if not text:
            return text
        
        import unicodedata
        
        # Remove excessive spaces first
        text = ' '.join(text.split())
        
        # COMPREHENSIVE Arabic Unicode artifact cleaning (expanded mapping)
        artifacts_map = {
            # Hamza forms
            'ﺀ': 'ء', 'ﺁ': 'آ', 'ﺂ': 'آ', 
            # Alef forms
            'ﺍ': 'ا', 'ﺎ': 'ا', 'ﺃ': 'أ', 'ﺄ': 'أ', 'ﺇ': 'إ', 'ﺈ': 'إ', 'ﺁ': 'آ', 'ﺂ': 'آ',
            # Alef with hamza
            'ﺅ': 'ؤ', 'ﺆ': 'ؤ', 'ﺋ': 'ئ', 'ﺌ': 'ئ', 'ﺉ': 'ئ', 'ﺊ': 'ئ',
            # Ba forms  
            'ﺏ': 'ب', 'ﺐ': 'ب', 'ﺑ': 'ب', 'ﺒ': 'ب',
            # Ta forms
            'ﺕ': 'ت', 'ﺖ': 'ت', 'ﺗ': 'ت', 'ﺘ': 'ت', 'ﺓ': 'ة', 'ﺔ': 'ة',
            # Tha forms
            'ﺙ': 'ث', 'ﺚ': 'ث', 'ﺛ': 'ث', 'ﺜ': 'ث',
            # Jeem forms
            'ﺝ': 'ج', 'ﺞ': 'ج', 'ﺟ': 'ج', 'ﺠ': 'ج',
            # Hha forms
            'ﺡ': 'ح', 'ﺢ': 'ح', 'ﺣ': 'ح', 'ﺤ': 'ح',
            # Kha forms
            'ﺥ': 'خ', 'ﺦ': 'خ', 'ﺧ': 'خ', 'ﺨ': 'خ',
            # Dal forms
            'ﺩ': 'د', 'ﺪ': 'د',
            # Thal forms
            'ﺫ': 'ذ', 'ﺬ': 'ذ',
            # Ra forms
            'ﺭ': 'ر', 'ﺮ': 'ر',
            # Zain forms
            'ﺯ': 'ز', 'ﺰ': 'ز',
            # Seen forms
            'ﺱ': 'س', 'ﺲ': 'س', 'ﺳ': 'س', 'ﺴ': 'س',
            # Sheen forms
            'ﺵ': 'ش', 'ﺶ': 'ش', 'ﺷ': 'ش', 'ﺸ': 'ش',
            # Sad forms
            'ﺹ': 'ص', 'ﺺ': 'ص', 'ﺻ': 'ص', 'ﺼ': 'ص',
            # Dad forms
            'ﺽ': 'ض', 'ﺾ': 'ض', 'ﺿ': 'ض', 'ﻀ': 'ض',
            # Taa forms
            'ﻁ': 'ط', 'ﻂ': 'ط', 'ﻃ': 'ط', 'ﻄ': 'ط',
            # Dhaa forms
            'ﻅ': 'ظ', 'ﻆ': 'ظ', 'ﻇ': 'ظ', 'ﻈ': 'ظ',
            # Ain forms
            'ﻉ': 'ع', 'ﻊ': 'ع', 'ﻋ': 'ع', 'ﻌ': 'ع',
            # Ghain forms
            'ﻍ': 'غ', 'ﻎ': 'غ', 'ﻏ': 'غ', 'ﻐ': 'غ',
            # Fa forms
            'ﻑ': 'ف', 'ﻒ': 'ف', 'ﻓ': 'ف', 'ﻔ': 'ف',
            # Qaf forms
            'ﻕ': 'ق', 'ﻖ': 'ق', 'ﻗ': 'ق', 'ﻘ': 'ق',
            # Kaf forms
            'ﻙ': 'ك', 'ﻚ': 'ك', 'ﻛ': 'ك', 'ﻜ': 'ك',
            # Lam forms
            'ﻝ': 'ل', 'ﻞ': 'ل', 'ﻟ': 'ل', 'ﻠ': 'ل',
            # Meem forms
            'ﻡ': 'م', 'ﻢ': 'م', 'ﻣ': 'م', 'ﻤ': 'م',
            # Noon forms
            'ﻥ': 'ن', 'ﻦ': 'ن', 'ﻧ': 'ن', 'ﻨ': 'ن',
            # Ha forms
            'ﻩ': 'ه', 'ﻪ': 'ه', 'ﻫ': 'ه', 'ﻬ': 'ه',
            # Waw forms
            'ﻭ': 'و', 'ﻮ': 'و',
            # Ya forms
            'ﻱ': 'ي', 'ﻲ': 'ي', 'ﻳ': 'ي', 'ﻴ': 'ي',
            # Lam-Alef ligatures
            'ﻼ': 'لا', 'ﻻ': 'لا', 'ﻷ': 'لأ', 'ﻹ': 'لإ', 'ﻵ': 'لآ',
        }
        
        # Apply character mapping
        for artifact, correct_char in artifacts_map.items():
            text = text.replace(artifact, correct_char)
        
        # Final cleanup: normalize Unicode (NFC form)
        try:
            text = unicodedata.normalize('NFC', text)
        except:
            pass
        
        return text

    def _normalize_fragmented_arabic(self, text: str) -> str:
        """Merge fragmented Arabic letters back into words"""
        if not text.strip():
            return text
        
        words = text.split()
        current_word = ""
        normalized_words = []
        
        for word in words:
            word_clean = word.strip()
            
            if not word_clean:
                continue
                
            # If this is a single Arabic character
            if len(word_clean) == 1 and '\u0600' <= word_clean <= '\u06FF':
                # Merge with current_word if it's building an Arabic word
                if current_word and '\u0600' <= current_word[-1] <= '\u06FF':
                    current_word += word_clean
                else:
                    if current_word:
                        normalized_words.append(current_word)
                    current_word = word_clean
            
            # If this is a number or English, separate it
            elif word_clean.isdigit() or word_clean.isalpha() or word_clean in ['.', ',', ':', ';']:
                if current_word:
                    normalized_words.append(current_word)
                    current_word = ""
                normalized_words.append(word_clean)
            
            # If this is a longer word that contains Arabic
            else:
                arabic_chars = sum(1 for c in word_clean if '\u0600' <= c <= '\u06FF')
                if arabic_chars > len(word_clean) * 0.7:  # Mostly Arabic word
                    if current_word and '\u0600' <= current_word[-1] <= '\u06FF':
                        current_word += word_clean
                    else:
                        if current_word:
                            normalized_words.append(current_word)
                            current_word = ""
                        current_word = word_clean
                else:  # Non-Arabic word
                    if current_word:
                        normalized_words.append(current_word)
                        current_word = ""
                    normalized_words.append(word_clean)
        
        # Add any remaining word
        if current_word:
            normalized_words.append(current_word)
        
        # Clean up artifacts and excessive spaces
        normalized_text = ' '.join(normalized_words)
        normalized_text = self._clean_text_artifacts(normalized_text)
        
        return normalized_text

    def _fix_arabic_text(self, text: str) -> str:
        """Comprehensive Arabic text fixing with proper RTL handling"""
        try:
            import arabic_reshaper
            from bidi.algorithm import get_display
        except ImportError:
            logger.warning("Arabic text processing libraries not available")
            return text
        
        if not text.strip():
            return text
        
        # Step 1: Clean Unicode artifacts first
        cleaned_text = self._clean_text_artifacts(text)
        
        # Step 2: Normalize fragmented text - merge broken letters into words
        normalized = self._normalize_fragmented_arabic(cleaned_text)
        
        # Step 3: Apply reshaping to Arabic words
        words = normalized.split()
        fixed_words = []
        
        for word in words:
            # Check if word contains Arabic characters
            arabic_chars = sum(1 for c in word if '\u0600' <= c <= '\u06FF')
            if arabic_chars > 0:
                # Apply reshaping to connect Arabic letters properly
                reshaped_word = arabic_reshaper.reshape(word)
                fixed_words.append(reshaped_word)
            else:
                # Keep non-Arabic words as-is
                fixed_words.append(word)
        
        # Step 4: Join words and apply BiDi algorithm (simple, no extra marks)
        text_for_bidi = ' '.join(fixed_words)
        
        # Apply BiDi algorithm simply
        try:
            fixed_text = get_display(text_for_bidi)
        except Exception as e:
            logger.warning(f"BiDi processing error: {e}")
            fixed_text = text_for_bidi
        
        return fixed_text

    def _ensure_rtl_text_direction(self, text: str) -> str:
        """Ensure Arabic text is displayed in proper RTL direction"""
        if not text.strip():
            return text
        
        # Check if text contains Arabic characters
        arabic_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF')
        if arabic_chars == 0:
            return text
        
        # Split text into lines and process each line
        lines = text.split('\n')
        rtl_lines = []
        
        for line in lines:
            if line.strip() and arabic_chars > 0:
                # Split line into words and process each word
                words = line.split()
                processed_words = []
                
                for word in words:
                    # Check if word contains Arabic
                    word_arabic_chars = sum(1 for c in word if '\u0600' <= c <= '\u06FF')
                    if word_arabic_chars > 0:
                        # For Arabic words, apply RTL mark
                        processed_word = '\u200F' + word + '\u200F'
                        processed_words.append(processed_word)
                    else:
                        processed_words.append(word)
                
                # Rejoin words and add RTL mark to the entire line
                processed_line = '\u202E' + ' '.join(processed_words) + '\u202C'
                rtl_lines.append(processed_line)
            else:
                rtl_lines.append(line)
        
        return '\n'.join(rtl_lines)
    
    # =====================================================
    # SECTION SEGMENTATION
    # =====================================================
    
    def split_case_sections(self, text: str) -> Dict[str, str]:
        """
        Split case text into logical sections based on Arabic keywords.
        
        Args:
            text: Full text of the legal case
            
        Returns:
            Dictionary with section types as keys and content as values
        """
        sections = {
            'summary': '',
            'facts': '',
            'arguments': '',
            'ruling': '',
            'legal_basis': ''
        }
        
        # Find all section markers with their positions
        section_markers = []
        
        for section_type, patterns in self.section_patterns.items():
            for pattern in patterns:
                for match in re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE):
                    section_markers.append({
                        'type': section_type,
                        'start': match.start(),
                        'end': match.end(),
                        'pattern': pattern
                    })
        
        # Sort markers by position
        section_markers.sort(key=lambda x: x['start'])
        
        # Extract content between markers
        for i, marker in enumerate(section_markers):
            section_type = marker['type']
            start = marker['end']
            
            # Find end position (start of next marker or end of text)
            if i + 1 < len(section_markers):
                end = section_markers[i + 1]['start']
            else:
                end = len(text)
            
            content = text[start:end].strip()
            
            # Append to existing content if section already found
            if sections[section_type]:
                sections[section_type] += "\n\n" + content
            else:
                sections[section_type] = content
        
        # If no sections found, put everything in summary
        if not any(sections.values()):
            logger.warning("No section markers found, using entire text as summary")
            sections['summary'] = text
        
        # Log what was found
        found_sections = [k for k, v in sections.items() if v]
        logger.info(f"Found sections: {', '.join(found_sections)}")
        
        return sections
    
    # =====================================================
    # DATABASE STORAGE
    # =====================================================
    
    async def save_case_with_sections(
        self,
        case_metadata: Dict[str, Any],
        sections: Dict[str, str],
        document_id: int
    ) -> LegalCase:
        """
        Save LegalCase and CaseSection records to database.
        
        Args:
            case_metadata: Dictionary containing case metadata
            sections: Dictionary of section content
            document_id: ID of the KnowledgeDocument
            
        Returns:
            Created LegalCase instance
        """
        # Parse date if provided as string
        decision_date = case_metadata.get('decision_date')
        if isinstance(decision_date, str):
            try:
                decision_date = datetime.strptime(decision_date, '%Y-%m-%d').date()
            except:
                decision_date = None
        
        # Create LegalCase record
        legal_case = LegalCase(
            case_number=case_metadata.get('case_number'),
            title=case_metadata.get('title'),
            description=case_metadata.get('description'),
            jurisdiction=case_metadata.get('jurisdiction'),
            court_name=case_metadata.get('court_name'),
            decision_date=decision_date,
            document_id=document_id,
            case_type=case_metadata.get('case_type'),
            court_level=case_metadata.get('court_level'),
            status='raw',
            created_at=datetime.utcnow()
        )
        
        self.db.add(legal_case)
        await self.db.flush()  # Get the ID
        
        logger.info(f"Created LegalCase ID: {legal_case.id}")
        
        # Create CaseSection records
        section_count = 0
        for section_type, content in sections.items():
            if content and content.strip():
                case_section = CaseSection(
                    case_id=legal_case.id,
                    section_type=section_type,
                    content=content.strip(),
                    created_at=datetime.utcnow()
                )
                self.db.add(case_section)
                section_count += 1
        
        await self.db.flush()
        
        logger.info(f"Created {section_count} CaseSection records")
        
        return legal_case
    
    # =====================================================
    # COMPLETE INGESTION PIPELINE
    # =====================================================
    
    async def ingest_legal_case(
        self,
        file_content: bytes,
        filename: str,
        case_metadata: Dict[str, Any],
        uploaded_by: int
    ) -> Dict[str, Any]:
        """
        Complete pipeline: upload file, extract text, segment sections, save to DB.
        
        Args:
            file_content: Binary content of the uploaded file
            filename: Original filename
            case_metadata: Dictionary containing case metadata
            uploaded_by: User ID who uploaded the file
            
        Returns:
            Dictionary with ingestion results
        """
        try:
            # Step 1: Save file and create KnowledgeDocument
            logger.info(f"Step 1: Saving uploaded file: {filename}")
            file_path, file_hash, knowledge_doc = await self.save_uploaded_case_file(
                file_content, filename, uploaded_by
            )
            
            # Step 2: Extract text from file
            logger.info(f"Step 2: Extracting text from: {file_path}")
            text = self.extract_text(file_path)
            
            if not text or len(text) < 50:
                raise ValueError(
                    f"Extracted text is too short ({len(text)} chars). "
                    "Possible causes:\n"
                    "  1. PDF is image-based (scanned) - Install Tesseract OCR for extraction\n"
                    "  2. PDF is mostly images with minimal text\n"
                    "  3. File is corrupted or password-protected\n"
                    "  4. Text extraction failed - check logs for details"
                )
            
            logger.info(f"Extracted {len(text)} characters")
            
            # Step 3: Split into sections
            logger.info(f"Step 3: Splitting text into sections")
            sections = self.split_case_sections(text)
            
            # Step 4: Save case and sections to database
            logger.info(f"Step 4: Saving case and sections to database")
            legal_case = await self.save_case_with_sections(
                case_metadata, sections, knowledge_doc.id
            )
            
            # Update KnowledgeDocument status
            knowledge_doc.status = 'processed'
            knowledge_doc.processed_at = datetime.utcnow()
            
            # Update LegalCase status
            legal_case.status = 'processed'
            
            # Commit transaction
            await self.db.commit()
            
            logger.info(f"✅ Successfully ingested legal case ID: {legal_case.id}")
            
            # Return results
            return {
                'success': True,
                'message': 'Legal case uploaded successfully',
                'data': {
                    'knowledge_document_id': knowledge_doc.id,
                    'legal_case_id': legal_case.id,
                    'case_number': legal_case.case_number,
                    'title': legal_case.title,
                    'file_path': file_path,
                    'file_hash': file_hash,
                    'text_length': len(text),
                    'sections_found': [k for k, v in sections.items() if v],
                    'sections_count': sum(1 for v in sections.values() if v)
                }
            }
            
        except Exception as e:
            await self.db.rollback()
            logger.error(f"Failed to ingest legal case: {str(e)}")
            
            # Clean up file if it was saved
            try:
                if 'file_path' in locals():
                    os.remove(file_path)
            except:
                pass
            
            return {
                'success': False,
                'message': f'Failed to ingest legal case: {str(e)}',
                'data': None
            }
    
    # =====================================================
    # BATCH INGESTION
    # =====================================================
    
    async def ingest_multiple_cases(
        self,
        cases_data: List[Dict[str, Any]],
        uploaded_by: int
    ) -> Dict[str, Any]:
        """
        Ingest multiple legal cases in batch.
        
        Args:
            cases_data: List of dictionaries, each containing:
                - file_content: bytes
                - filename: str
                - case_metadata: dict
            uploaded_by: User ID who uploaded the files
            
        Returns:
            Dictionary with batch ingestion results
        """
        results = {
            'total': len(cases_data),
            'successful': 0,
            'failed': 0,
            'cases': []
        }
        
        for i, case_data in enumerate(cases_data, 1):
            logger.info(f"Processing case {i}/{len(cases_data)}")
            
            result = await self.ingest_legal_case(
                file_content=case_data['file_content'],
                filename=case_data['filename'],
                case_metadata=case_data['case_metadata'],
                uploaded_by=uploaded_by
            )
            
            if result['success']:
                results['successful'] += 1
            else:
                results['failed'] += 1
            
            results['cases'].append({
                'filename': case_data['filename'],
                'success': result['success'],
                'message': result['message'],
                'data': result['data']
            })
        
        return results

