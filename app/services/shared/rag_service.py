"""
RAG (Retrieval-Augmented Generation) Service for Law Documents

This service handles:
- Document chunking and embedding generation
- Semantic search using embeddings
- Context retrieval for AI analysis
- Working with simplified LawDocument and LawChunk models
"""

import logging
import json
import re
import os
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_, func, delete, update
from sqlalchemy.orm import selectinload

from ...models.documnets import LawDocument, LawChunk
from .embedding_service import EmbeddingService

logger = logging.getLogger(__name__)


class RAGService:
    """
    RAG Service for Law Document processing and retrieval.
    
    Simplified to work with LawDocument and LawChunk models only.
    """
    
    def __init__(self, db: AsyncSession, model_name: str = 'legal_optimized'):
        """
        Initialize RAG Service.
        
        Args:
            db: Async database session
            model_name: Embedding model to use
        """
        self.db = db
        self.embedding_service = EmbeddingService(db, model_name=model_name)
        
        # Chunking settings
        self.chunk_size = 500  # words
        self.chunk_overlap = 50  # words
        self.min_chunk_size = 100  # words
        
        logger.info(f"✅ RAG Service initialized with model: {model_name}")
    
    def _smart_chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Smart chunking for Arabic legal text - محسّن للعربية.
        
        Args:
            text: Input text to chunk
            
        Returns:
            List of chunk dictionaries with content and metadata
        """
        if not text or not text.strip():
            return []
        
        # تنظيف النص العربي
        text = self._clean_arabic_text(text)
        
        # تقسيم إلى فقرات أولاً
        paragraphs = re.split(r'\n\s*\n', text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        
        chunks = []
        chunk_index = 0
        
        for paragraph in paragraphs:
            words = paragraph.split()
            
            # إذا كانت الفقرة قصيرة، أضفها كما هي
            if len(words) <= self.min_chunk_size:
                if len(words) >= 20:  # تجنب الفقرات القصيرة جداً
                    chunks.append({
                        'content': paragraph,
                        'word_count': len(words),
                        'chunk_index': chunk_index
                    })
                    chunk_index += 1
                continue
            
            # إذا كانت الفقرة طويلة، تقسيمها
            if len(words) > self.chunk_size:
                # تقسيم الفقرة الطويلة
                i = 0
                while i < len(words):
                    end_idx = min(i + self.chunk_size, len(words))
                    chunk_words = words[i:end_idx]
                    
                    # محاولة العثور على نهاية جملة طبيعية
                    if end_idx < len(words):
                        # البحث عن نهاية جملة في آخر 10 كلمات
                        for j in range(len(chunk_words) - 1, max(0, len(chunk_words) - 10), -1):
                            if any(chunk_words[j].endswith(end) for end in ['.', '۔', '؟', '!', '،']):
                                chunk_words = chunk_words[:j + 1]
                                break
                    
                    if len(chunk_words) >= self.min_chunk_size // 2:  # تجنب chunks صغيرة جداً
                        chunk_text = ' '.join(chunk_words)
                        chunks.append({
                            'content': chunk_text,
                            'word_count': len(chunk_words),
                            'chunk_index': chunk_index
                        })
                        chunk_index += 1
                    
                    # التحرك للـ chunk التالي مع overlap
                    i += len(chunk_words) - self.chunk_overlap
            else:
                # الفقرة ضمن الحجم المقبول
                chunks.append({
                    'content': paragraph,
                    'word_count': len(words),
                    'chunk_index': chunk_index
                })
                chunk_index += 1
        
        return chunks
    
    async def process_document(
        self,
        document_id: int,
        text: str,
        generate_embeddings: bool = True
    ) -> Dict[str, Any]:
        """
        Process a law document: chunk it and optionally generate embeddings.
        
        Args:
            document_id: ID of the law document
            text: Document text content
            generate_embeddings: Whether to generate embeddings immediately
            
        Returns:
            Processing results dictionary
        """
        try:
            logger.info(f"📄 Processing document {document_id}")
            
            # Get document
            doc_query = select(LawDocument).where(LawDocument.id == document_id)
            doc_result = await self.db.execute(doc_query)
            document = doc_result.scalar_one_or_none()
            
            if not document:
                raise ValueError(f"Document {document_id} not found")
            
            # Update status
            document.status = 'processing'
            await self.db.commit()
            
            # Chunk the text
            chunks_data = self._smart_chunk_text(text)
            
            # Update document
            document.total_chunks = len(chunks_data)
            document.processed_chunks = 0
            
            # Delete old chunks
            await self.db.execute(
                delete(LawChunk).where(LawChunk.document_id == document_id)
            )
            
            # Create new chunks
            chunks_created = []
            for chunk_data in chunks_data:
                chunk = LawChunk(
                    document_id=document_id,
                    content=chunk_data['content'],
                    word_count=chunk_data['word_count'],
                    chunk_index=chunk_data['chunk_index'],
                    is_processed=False
                )
                self.db.add(chunk)
                chunks_created.append(chunk)
            
            await self.db.commit()
            
            logger.info(f"✅ Created {len(chunks_created)} chunks for document {document_id}")
            
            # Generate embeddings if requested
            if generate_embeddings:
                logger.info(f"🔄 Generating embeddings for document {document_id}")
                await self.generate_embeddings_for_document(document_id)
            
            # Update document status
            document.status = 'completed' if generate_embeddings else 'chunked'
            document.processed_at = datetime.utcnow()
            await self.db.commit()
            
            return {
                'success': True,
                'document_id': document_id,
                'total_chunks': len(chunks_created),
                'embeddings_generated': generate_embeddings,
                'status': document.status
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing document {document_id}: {str(e)}")
            
            # Update document with error
            if document:
                document.status = 'failed'
                document.error_message = str(e)
                await self.db.commit()
            
            return {
                'success': False,
                'document_id': document_id,
                'error': str(e)
            }
    
    async def generate_embeddings_for_document(
        self,
        document_id: int,
        batch_size: int = 16
    ) -> Dict[str, Any]:
        """
        Generate embeddings for all chunks of a document.
        
        Args:
            document_id: ID of the document
            batch_size: Number of chunks to process at once
            
        Returns:
            Results dictionary
        """
        try:
            logger.info(f"🔄 Generating embeddings for document {document_id}")
            
            # Get document and chunks
            doc_query = select(LawDocument).where(LawDocument.id == document_id)
            doc_result = await self.db.execute(doc_query)
            document = doc_result.scalar_one_or_none()
            
            if not document:
                raise ValueError(f"Document {document_id} not found")
            
            # Get unprocessed chunks
            chunks_query = (
                select(LawChunk)
                .where(
                    and_(
                        LawChunk.document_id == document_id,
                        LawChunk.is_processed == False
                    )
                )
                .order_by(LawChunk.chunk_index)
            )
            chunks_result = await self.db.execute(chunks_query)
            chunks = chunks_result.scalars().all()
            
            if not chunks:
                logger.info(f"✅ No unprocessed chunks for document {document_id}")
                return {
                    'success': True,
                    'document_id': document_id,
                    'processed': 0
                }
            
            # Initialize embedding service
            await self.embedding_service.initialize()
            
            # Process chunks in batches
            processed_count = 0
            failed_count = 0
            
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                texts = [chunk.content for chunk in batch]
                
                try:
                    # Generate embeddings
                    embeddings = await self.embedding_service.generate_embeddings_batch(texts)
                    
                    # Update chunks
                    for chunk, embedding in zip(batch, embeddings):
                        chunk.embedding_vector = json.dumps(embedding.tolist())
                        chunk.is_processed = True
                        processed_count += 1
                    
                    # Update document progress
                    document.processed_chunks = processed_count
                    await self.db.commit()
                    
                    logger.info(f"✅ Processed batch {i // batch_size + 1}, total: {processed_count}/{len(chunks)}")
                    
                except Exception as e:
                    logger.error(f"❌ Error processing batch: {str(e)}")
                    failed_count += len(batch)
            
            # Final update
            document.status = 'completed' if failed_count == 0 else 'partially_completed'
            document.processed_at = datetime.utcnow()
            await self.db.commit()
            
            logger.info(f"✅ Embeddings generation complete for document {document_id}: {processed_count} success, {failed_count} failed")
            
            return {
                'success': True,
                'document_id': document_id,
                'processed': processed_count,
                'failed': failed_count,
                'total': len(chunks)
            }
            
        except Exception as e:
            logger.error(f"❌ Error generating embeddings for document {document_id}: {str(e)}")
            return {
                'success': False,
                'document_id': document_id,
                'error': str(e)
            }
    
    async def semantic_search(
        self,
        query: str,
        top_k: int = 10,
        threshold: float = 0.7,
        document_id: Optional[int] = None,
        jurisdiction: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        Semantic search across law chunks.
        
        Args:
            query: Search query
            top_k: Number of results to return
            threshold: Minimum similarity threshold
            document_id: Optional filter by specific document
            jurisdiction: Optional filter by jurisdiction
            
        Returns:
            List of search results with similarity scores
        """
        try:
            logger.info(f"🔍 Semantic search: '{query[:50]}...'")
            
            # Initialize embedding service
            await self.embedding_service.initialize()
            
            # Generate query embedding
            query_embedding = await self.embedding_service.generate_embedding(query)
            
            # Build query
            query_builder = (
                select(LawChunk, LawDocument)
                .join(LawDocument, LawChunk.document_id == LawDocument.id)
                .where(
                    and_(
                        LawChunk.embedding_vector.isnot(None),
                        LawChunk.embedding_vector != '',
                        LawChunk.is_processed == True
                    )
                )
            )
            
            # Apply filters
            if document_id:
                query_builder = query_builder.where(LawChunk.document_id == document_id)
            
            if jurisdiction:
                query_builder = query_builder.where(LawDocument.jurisdiction.ilike(f"%{jurisdiction}%"))
            
            # Execute query
            result = await self.db.execute(query_builder)
            rows = result.all()
            
            logger.info(f"📊 Found {len(rows)} chunks to search")
            
            # Calculate similarities
            results = []
            for chunk, document in rows:
                try:
                    chunk_embedding = json.loads(chunk.embedding_vector)
                    similarity = self.embedding_service.cosine_similarity(
                        query_embedding,
                        chunk_embedding
                    )
                    
                    if similarity >= threshold:
                        results.append({
                            'chunk_id': chunk.id,
                            'document_id': document.id,
                            'document_name': document.name,
                            'jurisdiction': document.jurisdiction,
                            'content': chunk.content,
                            'similarity': float(similarity),
                            'chunk_index': chunk.chunk_index,
                            'word_count': chunk.word_count
                        })
                        
                except Exception as e:
                    logger.warning(f"⚠️  Error processing chunk {chunk.id}: {str(e)}")
                    continue
            
            # Sort by similarity
            results.sort(key=lambda x: x['similarity'], reverse=True)
            results = results[:top_k]
            
            logger.info(f"✅ Found {len(results)} results above threshold")
            
            return results
            
        except Exception as e:
            logger.error(f"❌ Semantic search failed: {str(e)}")
            return []
    
    async def get_context_for_query(
        self,
        query: str,
        top_k: int = 5,
        max_context_length: int = 2000
    ) -> str:
        """
        Get relevant context for a query (for RAG).
        
        Args:
            query: User query
            top_k: Number of chunks to retrieve
            max_context_length: Maximum total context length in words
            
        Returns:
            Combined context string
        """
        # Get relevant chunks
        results = await self.semantic_search(
            query=query,
            top_k=top_k,
            threshold=0.6
        )
        
        if not results:
            return ""
        
        # Build context
        context_parts = []
        total_words = 0
        
        for result in results:
            content = result['content']
            words = len(content.split())
            
            if total_words + words > max_context_length:
                break
            
            context_parts.append(f"[المصدر: {result['document_name']}]\n{content}")
            total_words += words
        
        return "\n\n---\n\n".join(context_parts)
    
    async def get_statistics(self) -> Dict[str, Any]:
        """
        Get RAG service statistics.
        
        Returns:
            Statistics dictionary
        """
        try:
            # Total documents
            total_docs = await self.db.execute(select(func.count(LawDocument.id)))
            total_documents = total_docs.scalar() or 0
            
            # Total chunks
            total_chunks_result = await self.db.execute(select(func.count(LawChunk.id)))
            total_chunks = total_chunks_result.scalar() or 0
            
            # Chunks with embeddings
            with_embeddings = await self.db.execute(
                select(func.count(LawChunk.id)).where(
                    and_(
                        LawChunk.embedding_vector.isnot(None),
                        LawChunk.embedding_vector != '',
                        LawChunk.is_processed == True
                    )
                )
            )
            chunks_with_embeddings = with_embeddings.scalar() or 0
            
            # Documents by status
            status_counts = await self.db.execute(
                select(
                    LawDocument.status,
                    func.count(LawDocument.id)
                ).group_by(LawDocument.status)
            )
            status_breakdown = {status: count for status, count in status_counts}
            
            return {
                'total_documents': total_documents,
                'total_chunks': total_chunks,
                'chunks_with_embeddings': chunks_with_embeddings,
                'embedding_coverage': f"{(chunks_with_embeddings / total_chunks * 100):.1f}%" if total_chunks > 0 else "0%",
                'documents_by_status': status_breakdown,
                'model': self.embedding_service.model_name
            }
            
        except Exception as e:
            logger.error(f"❌ Error getting statistics: {str(e)}")
            return {
                'error': str(e)
            }
    
    def _clean_arabic_text(self, text: str) -> str:
        """
        تنظيف النص العربي.
        
        Args:
            text: النص المراد تنظيفه
            
        Returns:
            النص المنظف
        """
        if not text:
            return ""
        
        # إزالة المسافات الزائدة
        text = re.sub(r'\s+', ' ', text)
        
        # إزالة الترقيم الزائد مع الحفاظ على نهايات الجمل
        text = re.sub(r'[^\w\u0600-\u06FF\s\.\?\!،؛]', ' ', text)
        
        return text.strip()
    
    async def _read_document_file(self, file_path: str) -> Dict:
        """
        قراءة ملفات بأنواع مختلفة.
        
        Args:
            file_path: مسار الملف
            
        Returns:
            قاموس يحتوي على النص والمعلومات
        """
        try:
            file_size = os.path.getsize(file_path)
            
            if file_path.endswith('.docx'):
                try:
                    import docx
                    doc = docx.Document(file_path)
                    full_text = "\n".join([
                        paragraph.text for paragraph in doc.paragraphs 
                        if paragraph.text.strip()
                    ])
                except ImportError:
                    raise ImportError("Please install python-docx: pip install python-docx")
                    
            elif file_path.endswith('.pdf'):
                try:
                    import PyPDF2
                    full_text = ""
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        for page in reader.pages:
                            text = page.extract_text()
                            if text:
                                full_text += text + "\n"
                except ImportError:
                    raise ImportError("Please install PyPDF2: pip install PyPDF2")
                            
            elif file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as file:
                    full_text = file.read()
                    
            else:
                raise ValueError(f"Unsupported file format: {file_path}")
            
            return {
                'full_text': full_text,
                'file_type': file_path.split('.')[-1].upper(),
                'file_size': file_size
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to read document: {e}")
            raise
    
    async def _get_document_chunks(self, document_id: int) -> List[LawChunk]:
        """
        الحصول على chunks المستند.
        
        Args:
            document_id: معرف المستند
            
        Returns:
            قائمة بـ chunks المستند
        """
        query = select(LawChunk).where(LawChunk.document_id == document_id)
        result = await self.db.execute(query)
        return result.scalars().all()
    
    async def ingest_law_document(self, file_path: str, law_metadata: Dict) -> Dict:
        """
        استيعاب قانون مباشرة من ملف - النسخة المحدثة.
        
        Args:
            file_path: مسار الملف
            law_metadata: بيانات وصفية عن القانون
            
        Returns:
            نتيجة المعالجة
        """
        try:
            logger.info(f"📥 Ingesting law document: {law_metadata.get('law_name', 'Unknown')}")
            
            # 1. قراءة الملف
            document_data = await self._read_document_file(file_path)
            
            # 2. إنشاء مستند في قاعدة البيانات
            law_doc = LawDocument(
                name=law_metadata.get('law_name', 'Unnamed Document'),
                type=law_metadata.get('law_type', 'law'),
                jurisdiction=law_metadata.get('jurisdiction', 'السعودية'),
                original_filename=law_metadata.get('original_filename', os.path.basename(file_path)),
                file_size=document_data.get('file_size', 0),
                status='processing'
            )
            self.db.add(law_doc)
            await self.db.commit()
            await self.db.refresh(law_doc)
            
            logger.info(f"✅ Created LawDocument: {law_doc.id}")
            
            # 3. معالجة النص وإنشاء chunks
            processing_result = await self.process_document(
                document_id=law_doc.id,
                text=document_data['full_text'],
                generate_embeddings=True
            )
            
            if processing_result['success']:
                # الحصول على chunks للإحصائيات
                chunks = await self._get_document_chunks(law_doc.id)
                total_words = sum(chunk.word_count or 0 for chunk in chunks)
                
                return {
                    'success': True,
                    'law_name': law_metadata.get('law_name', 'Unnamed'),
                    'chunks_created': processing_result['total_chunks'],
                    'chunks_stored': processing_result['total_chunks'],
                    'processing_time': 0.0,
                    'file_type': document_data.get('file_type', 'UNKNOWN'),
                    'total_words': total_words,
                    'document_id': law_doc.id
                }
            else:
                # تحديث حالة المستند إلى فاشل
                law_doc.status = 'failed'
                law_doc.error_message = processing_result.get('error', 'Unknown error')
                await self.db.commit()
                
                return {
                    'success': False,
                    'error': processing_result.get('error', 'Processing failed')
                }
                
        except Exception as e:
            logger.error(f"❌ Law document ingestion failed: {str(e)}", exc_info=True)
            return {
                'success': False,
                'error': str(e)
            }
    
    async def search(
        self,
        query: str,
        top_k: int = 5,
        threshold: float = 0.6,
        law_source_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        البحث الدلالي - متوافق مع واجهة API.
        
        Args:
            query: استعلام البحث
            top_k: عدد النتائج
            threshold: حد التشابه
            law_source_id: للحفاظ على التوافق (document_id)
            
        Returns:
            نتائج البحث بتنسيق API
        """
        start_time = datetime.utcnow()
        
        try:
            logger.info(f"🔍 RAG Search: '{query[:50]}...'")
            
            # تحضير الفلاتر
            filters = {}
            if law_source_id:
                filters['document_id'] = law_source_id
            
            # إجراء البحث
            results = await self.semantic_search(
                query=query,
                top_k=top_k,
                threshold=threshold,
                document_id=filters.get('document_id')
            )
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            # تنسيق النتائج لتتوافق مع الـ API
            formatted_results = []
            for result in results:
                formatted_results.append({
                    'chunk_id': result['chunk_id'],
                    'content': result['content'],
                    'similarity_score': round(result['similarity'], 4),
                    'law_source_id': result['document_id'],
                    'law_source_name': result['document_name'],
                    'word_count': result['word_count'],
                    'metadata': {
                        'jurisdiction': result.get('jurisdiction', ''),
                        'chunk_index': result.get('chunk_index', 0)
                    }
                })
            
            return {
                'success': True,
                'query': query,
                'total_results': len(formatted_results),
                'results': formatted_results,
                'processing_time': round(processing_time, 2)
            }
            
        except Exception as e:
            logger.error(f"❌ Search failed: {str(e)}")
            return {
                'success': False,
                'query': query,
                'total_results': 0,
                'results': [],
                'processing_time': 0.0,
                'error': str(e)
            }
    
    async def get_system_status(self) -> Dict[str, Any]:
        """
        الحصول على حالة النظام - مطلوبة من قبل API.
        
        Returns:
            حالة النظام
        """
        try:
            stats = await self.get_statistics()
            
            if 'error' in stats:
                return {'status': 'error', 'error': stats['error']}
            
            total_chunks = stats.get('total_chunks', 0)
            chunks_with_embeddings = stats.get('chunks_with_embeddings', 0)
            
            coverage = (chunks_with_embeddings / total_chunks * 100) if total_chunks > 0 else 0
            
            return {
                'status': 'operational',
                'total_documents': stats.get('total_documents', 0),
                'total_chunks': total_chunks,
                'chunks_with_embeddings': chunks_with_embeddings,
                'embedding_coverage': round(coverage, 2),
                'documents_by_status': stats.get('documents_by_status', {}),
                'chunking_settings': {
                    'max_chunk_words': self.chunk_size,
                    'min_chunk_words': self.min_chunk_size,
                    'chunk_overlap_words': self.chunk_overlap
                },
                'model': stats.get('model', 'unknown'),
                'timestamp': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to get system status: {str(e)}")
            return {
                'status': 'error',
                'error': str(e)
            }
