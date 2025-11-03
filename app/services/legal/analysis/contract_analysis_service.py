"""
Contract Analysis Service using Gemini AI.
This service analyzes contracts by sending files directly to Gemini AI.
"""

import os
import asyncio
import logging
from typing import Dict, Any, Optional
from google.genai import types

logger = logging.getLogger(__name__)


class ContractAnalysisService:
    """
    Service for analyzing contracts using Gemini AI.
    Sends files directly to Gemini without manual text extraction.
    """
    
    def __init__(self, api_key: Optional[str] = None):
        self.api_key = api_key or os.getenv("GEMINI_API_KEY")
        self._client = None
        self._initialize_client()
    
    def _initialize_client(self):
        """Initialize Gemini client if not already initialized."""
        if not self._client:
            try:
                from google import genai
                self._client = genai.Client(api_key=self.api_key)
                logger.info("Gemini client initialized successfully")
            except ImportError:
                logger.error("google-genai library not available")
                raise
            except Exception as e:
                logger.error(f"Failed to initialize Gemini client: {e}")
    
    async def analyze_contract(
        self,
        file_content: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """
        Analyze a contract file using Gemini AI.
        
        Args:
            file_content: The contract file content as bytes
            filename: Name of the contract file
            
        Returns:
            Dict containing analysis results with weak_points, risks, and suggestions
        """
        try:
            if not self._client:
                return {
                    "success": False,
                    "message": "Gemini client not initialized",
                    "data": None
                }
            
            # Validate file size
            file_size_mb = len(file_content) / (1024 * 1024)
            if file_size_mb > 20:
                return {
                    "success": False,
                    "message": f"File too large ({file_size_mb:.1f}MB). Maximum size is 20MB.",
                    "data": None
                }
            
            # Determine file extension and MIME type
            file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
            mime_type_map = {
                'pdf': 'application/pdf',
                'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'doc': 'application/msword',
                'txt': 'text/plain'
            }
            mime_type = mime_type_map.get(file_ext, 'application/octet-stream')
            
            # Create prompt for contract analysis
            prompt = self._create_contract_analysis_prompt()
            
            logger.info(f"Starting Gemini AI contract analysis for: {filename}")
            
            # Handle DOCX files - extract text first since Gemini doesn't support DOCX directly
            if file_ext in ['docx', 'doc']:
                # Extract text from DOCX/DOC file
                logger.info(f"Extracting text from {file_ext.upper()} file before sending to Gemini")
                extracted_text = await self._extract_text_from_docx(file_content, filename)
                
                if not extracted_text or not extracted_text.strip():
                    return {
                        "success": False,
                        "message": "Failed to extract text from DOCX/DOC file. The file may be corrupted or empty.",
                        "data": None
                    }
                
                logger.info(f"Extracted {len(extracted_text)} characters from {file_ext.upper()} file")
                
                # Send extracted text to Gemini
                content_parts = [
                    f"Contract Content from {filename}:\n\n{extracted_text}",
                    prompt
                ]
            else:
                # For PDF and TXT files, send directly to Gemini
                if file_ext == 'pdf':
                    file_part = types.Part.from_bytes(data=file_content, mime_type=mime_type)
                    content_parts = [file_part, prompt]
                elif file_ext == 'txt':
                    # For TXT files, decode and send as text
                    try:
                        text_content = file_content.decode('utf-8')
                    except UnicodeDecodeError:
                        try:
                            text_content = file_content.decode('utf-8-sig')  # Handle BOM
                        except UnicodeDecodeError:
                            text_content = file_content.decode('latin-1')  # Fallback
                    
                    content_parts = [
                        f"Contract Content from {filename}:\n\n{text_content}",
                        prompt
                    ]
                else:
                    return {
                        "success": False,
                        "message": f"Unsupported file type for direct processing: {file_ext}",
                        "data": None
                    }
            
            # Call Gemini API with timeout
            try:
                response = await asyncio.wait_for(
                    asyncio.to_thread(
                        self._client.models.generate_content,
                        model="gemini-2.0-flash-exp",
                        contents=content_parts
                    ),
                    timeout=300  # 5 minutes timeout
                )
                
                analysis_text = getattr(response, "text", "")
                if not analysis_text:
                    return {
                        "success": False,
                        "message": "Gemini AI returned empty response",
                        "data": None
                    }
                
                logger.info("Gemini AI contract analysis completed successfully")
                
                # Parse JSON response from Gemini
                analysis_data = self._parse_contract_analysis_response(analysis_text)
                
                return {
                    "success": True,
                    "message": "Contract analysis completed successfully",
                    "data": analysis_data
                }
                
            except asyncio.TimeoutError:
                logger.error("Gemini AI contract analysis timed out after 5 minutes")
                return {
                    "success": False,
                    "message": "Analysis timed out. Please try with a smaller file or try again later.",
                    "data": None
                }
            except Exception as e:
                logger.error(f"Gemini AI API call failed: {e}", exc_info=True)
                return {
                    "success": False,
                    "message": f"AI analysis failed: {str(e)}",
                    "data": None
                }
                
        except Exception as e:
            logger.error(f"Contract analysis error: {e}", exc_info=True)
            return {
                "success": False,
                "message": f"Failed to analyze contract: {str(e)}",
                "data": None
            }
    
    def _create_contract_analysis_prompt(self) -> str:
        """
        Create a prompt for contract analysis that instructs Gemini to return JSON.
        The prompt detects the contract language and responds in the same language.
        """
        return """You are a legal expert specializing in contract analysis. 

IMPORTANT INSTRUCTIONS:
1. First, detect the primary language of the contract document (Arabic or English)
2. Analyze the contract thoroughly
3. Return your ENTIRE response (including all text) in the SAME language as the contract
   - If the contract is in Arabic, respond in Arabic
   - If the contract is in English, respond in English
4. Return ONLY valid JSON, no additional text before or after

JSON Structure:
{
  "weak_points": ["list of weak points or problematic clauses"],
  "risks": ["list of potential legal and business risks"],
  "suggestions": ["list of suggested changes or new clauses to add"]
}

REQUIREMENTS:
- Each array should contain clear, specific, and actionable items in the contract's language
- Focus on legal, financial, and business risks
- Provide practical suggestions for improvement
- If a section is empty, use an empty array []
- All text must be in the same language as the contract

Example (if contract is in English):
{
  "weak_points": ["Vague termination clause", "Missing force majeure provision"],
  "risks": ["No dispute resolution mechanism", "Unclear payment terms"],
  "suggestions": ["Add arbitration clause", "Specify payment schedule", "Include indemnification clause"]
}

Example (if contract is in Arabic):
{
  "weak_points": ["شرط الإنهاء غامض", "عدم وجود بند القوة القاهرة"],
  "risks": ["لا يوجد آلية لحل النزاعات", "شروط الدفع غير واضحة"],
  "suggestions": ["إضافة شرط التحكيم", "تحديد جدول الدفع", "إضافة شرط التعويض"]
}

Now analyze the contract, detect its language, and return the JSON response in that language:"""
    
    def _parse_contract_analysis_response(self, text: str) -> Dict[str, Any]:
        """
        Parse the Gemini response to extract JSON structure.
        Returns a dictionary with weak_points, risks, and suggestions.
        """
        import json
        import re
        
        # Try to parse as direct JSON
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass
        
        # Try to extract JSON from code blocks
        try:
            json_match = re.search(r'```(?:json)?\s*(\{[\s\S]*?\})\s*```', text)
            if json_match:
                return json.loads(json_match.group(1))
        except (json.JSONDecodeError, AttributeError):
            pass
        
        # Try to find JSON object in the text
        try:
            start_idx = text.find('{')
            end_idx = text.rfind('}')
            if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
                json_str = text[start_idx:end_idx + 1]
                parsed = json.loads(json_str)
                # Ensure all required fields exist
                return {
                    "weak_points": parsed.get("weak_points", []),
                    "risks": parsed.get("risks", []),
                    "suggestions": parsed.get("suggestions", [])
                }
        except (json.JSONDecodeError, ValueError):
            pass
        
        # Fallback: if JSON parsing fails, try to extract information from text
        logger.warning("Failed to parse JSON from Gemini response, using fallback parsing")
        return {
            "weak_points": self._extract_list_from_text(text, "weak_points", "weak point"),
            "risks": self._extract_list_from_text(text, "risks", "risk"),
            "suggestions": self._extract_list_from_text(text, "suggestions", "suggestion")
        }
    
    def _extract_list_from_text(self, text: str, key: str, keyword: str) -> list:
        """Fallback: Extract list items from text when JSON parsing fails."""
        items = []
        lines = text.split('\n')
        capturing = False
        
        for line in lines:
            line_lower = line.lower()
            if key in line_lower or keyword in line_lower:
                capturing = True
            elif capturing and line.strip().startswith(('-', '•', '*', '1.', '2.', '3.')):
                item = line.strip().lstrip('-•*0123456789. ').strip()
                if item:
                    items.append(item)
            elif capturing and line.strip() and not line.strip().startswith(('{', '}', '[', ']')):
                item = line.strip()
                if item and len(item) > 10:  # Filter out very short items
                    items.append(item)
        
        return items[:10]  # Limit to 10 items
    
    async def _extract_text_from_docx(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOCX/DOC file."""
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document(BytesIO(file_content))
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n\n'.join(text_content)
        except ImportError:
            logger.error("python-docx not installed. Cannot extract text from DOCX files.")
            return ""
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""

