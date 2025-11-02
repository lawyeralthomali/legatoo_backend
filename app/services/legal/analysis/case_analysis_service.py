"""
Legal Case Analysis Service using Gemini AI.

This service provides comprehensive legal case analysis using Google Gemini AI,
tailored for Saudi Arabian law with detailed analysis for both lawyers and users.
"""

import os
import logging
import asyncio
import tempfile
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime

logger = logging.getLogger(__name__)

# Try to import python-docx for DOCX text extraction
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX files will need text extraction.")


class CaseAnalysisService:
    """
    Service for analyzing legal cases using Gemini AI.
    
    Provides detailed analysis according to Saudi Arabian law with comprehensive
    information suitable for both legal professionals and general users.
    """
    
    def __init__(self, api_key: Optional[str] = None):
        """Initialize the case analysis service."""
        self.api_key = api_key or os.getenv("GEMINI_API_KEY")
        self._client = None
        self._model = None
        
    def _initialize_client(self):
        """Initialize Gemini client if not already initialized."""
        if not self._client:
            try:
                from google import genai
                self._client = genai.Client(api_key=self.api_key)
                logger.info("Gemini client initialized successfully")
            except ImportError:
                logger.error("google-genai library not available")
                raise
            except Exception as e:
                logger.error(f"Failed to initialize Gemini client: {e}")
                raise
    
    async def analyze_case(
        self,
        file_content: bytes,
        filename: str,
        analysis_type: str,
        lawsuit_type: str,
        result_seeking: str,
        user_context: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Analyze a legal case file using Gemini AI.
        
        Args:
            file_content: Bytes content of the case file (PDF, DOCX, TXT)
            filename: Original filename
            analysis_type: Type of analysis ("case-analysis" or "contract-review")
            lawsuit_type: Type of lawsuit (e.g., "commercial", "labor", "civil")
            result_seeking: What the user wants to achieve
            user_context: Additional user context
            
        Returns:
            Dictionary with analysis results
        """
        try:
            # Initialize client
            self._initialize_client()
            
            # Determine MIME type
            file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
            mime_types = {
                'pdf': 'application/pdf',
                'doc': 'application/msword',
                'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'txt': 'text/plain'
            }
            mime_type = mime_types.get(file_ext)
            
            if not mime_type:
                return {
                    "success": False,
                    "message": f"Unsupported file type: {file_ext}",
                    "data": None
                }
            
            # Check file size (limit to 20MB)
            file_size_mb = len(file_content) / (1024 * 1024)
            if file_size_mb > 20:
                return {
                    "success": False,
                    "message": f"File too large ({file_size_mb:.1f}MB). Maximum size is 20MB.",
                    "data": None
                }
            
            # Create comprehensive analysis prompt for Saudi law
            prompt = self._create_analysis_prompt(
                analysis_type=analysis_type,
                lawsuit_type=lawsuit_type,
                result_seeking=result_seeking,
                user_context=user_context,
                filename=filename
            )
            
            logger.info(f"Starting Gemini AI analysis for: {filename}")
            logger.info(f"Analysis type: {analysis_type}, Lawsuit type: {lawsuit_type}")
            
            # Handle DOCX files differently - extract text first since Gemini doesn't support DOCX directly
            from google.genai import types
            
            if file_ext in ['docx', 'doc']:
                # Extract text from DOCX/DOC file
                logger.info(f"Extracting text from {file_ext.upper()} file before sending to Gemini")
                extracted_text = await self._extract_text_from_docx(file_content, filename)
                
                if not extracted_text or not extracted_text.strip():
                    return {
                        "success": False,
                        "message": "Failed to extract text from DOCX/DOC file. The file may be corrupted or empty.",
                        "data": None
                    }
                
                logger.info(f"Extracted {len(extracted_text)} characters from {file_ext.upper()} file")
                
                # Send extracted text to Gemini instead of binary file
                content_parts = [
                    f"Document Content from {filename}:\n\n{extracted_text}",
                    prompt
                ]
                
            else:
                # For PDF and TXT files, send directly to Gemini
                if file_ext == 'pdf':
                    file_part = types.Part.from_bytes(data=file_content, mime_type=mime_type)
                    content_parts = [file_part, prompt]
                elif file_ext == 'txt':
                    # For TXT files, decode and send as text
                    try:
                        text_content = file_content.decode('utf-8')
                    except UnicodeDecodeError:
                        try:
                            text_content = file_content.decode('utf-8-sig')  # Handle BOM
                        except UnicodeDecodeError:
                            text_content = file_content.decode('latin-1')  # Fallback
                    
                    content_parts = [
                        f"Document Content from {filename}:\n\n{text_content}",
                        prompt
                    ]
                else:
                    return {
                        "success": False,
                        "message": f"Unsupported file type for direct processing: {file_ext}",
                        "data": None
                    }
            
            # Call Gemini API with timeout
            try:
                response = await asyncio.wait_for(
                    asyncio.to_thread(
                        self._client.models.generate_content,
                        model="gemini-2.0-flash-exp",
                        contents=content_parts
                    ),
                    timeout=300  # 5 minutes timeout
                )
                
                analysis_text = getattr(response, "text", "")
                if not analysis_text:
                    return {
                        "success": False,
                        "message": "Gemini AI returned empty response",
                        "data": None
                    }
                
                logger.info("Gemini AI analysis completed successfully")
                
                # Parse and structure the analysis
                analysis_data = self._parse_analysis_response(analysis_text)
                
                return {
                    "success": True,
                    "message": "Case analysis completed successfully",
                    "data": {
                        "analysis_id": None,  # Can be stored in DB if needed
                        "filename": filename,
                        "uploaded_at": datetime.utcnow().isoformat(),
                        "analysis_type": analysis_type,
                        "lawsuit_type": lawsuit_type,
                        "result_seeking": result_seeking,
                        "analysis": analysis_data,
                        "raw_response": analysis_text,  # Include raw for debugging
                        "status": "completed"
                    }
                }
                
            except asyncio.TimeoutError:
                logger.error("Gemini AI analysis timed out after 5 minutes")
                return {
                    "success": False,
                    "message": "Analysis timed out. Please try with a smaller file or try again later.",
                    "data": None
                }
            except Exception as e:
                logger.error(f"Gemini AI API call failed: {e}", exc_info=True)
                return {
                    "success": False,
                    "message": f"AI analysis failed: {str(e)}",
                    "data": None
                }
                
        except Exception as e:
            logger.error(f"Case analysis error: {e}", exc_info=True)
            return {
                "success": False,
                "message": f"Failed to analyze case: {str(e)}",
                "data": None
            }
    
    def _create_analysis_prompt(
        self,
        analysis_type: str,
        lawsuit_type: str,
        result_seeking: str,
        user_context: Optional[str],
        filename: str
    ) -> str:
        """
        Create a comprehensive analysis prompt for Saudi law.
        
        Returns a detailed prompt that instructs Gemini to provide thorough
        analysis suitable for both lawyers and users.
        """
        
        lawsuit_type_ar = {
            "commercial": "تجاري",
            "labor": "عمل",
            "personal-status": "أحوال شخصية",
            "criminal": "جنائي",
            "civil": "مدني",
            "administrative": "إداري",
            "contract-dispute": "نزاع تعاقدي",
            "real-estate": "عقاري",
            "intellectual-property": "ملكية فكرية",
            "family": "أسري",
            "employment": "توظيف",
            "other": "أخرى"
        }.get(lawsuit_type, lawsuit_type)
        
        prompt = f"""
أنت خبير قانوني سعودي متخصص في النظام القانوني السعودي. مهمتك تحليل هذا المستند القانوني تحليلاً شاملاً ومفصلاً وفقاً للقوانين والأنظمة السعودية.

**نوع التحليل المطلوب:** {analysis_type}
**نوع القضية:** {lawsuit_type_ar} ({lawsuit_type})
**الهدف من التحليل:** {result_seeking}
{'**سياق إضافي:** ' + user_context if user_context else ''}

**المستند المراد تحليله:** {filename}

---

**ملاحظة مهمة:** 
- قد يحتوي المستند على قضية مكتملة (مع حكم) أو قد يكون طلب استشارة لتحليل نقاط القوة والضعف
- يجب تحليل المستند حسب محتواه الفعلي
- إذا كان المستند يحتوي على حكم قضائي، حلله. إذا كان طلب استشارة، ركز على النقاط القانونية الضعيفة والتوصيات

**التعليمات التفصيلية للتحليل:**

يجب أن يتضمن تحليلك العناصر التالية بالتفصيل الكامل - استخدم العناوين التالية بالضبط:

### 1. ملخص تنفيذي شامل (Executive Summary)
- ملخص شامل للقضية أو المستند
- النقاط الرئيسية والأطراف المعنية
- الطبيعة القانونية للنزاع أو الموضوع

### 2. التحليل القانوني المفصل (Detailed Legal Analysis)

#### أ. الوضع القانوني الحالي (Legal Status)
- تحليل الوضع القانوني الحالي للقضية أو الطلب
- الحقوق والالتزامات القانونية
- الوضع القانوني للأطراف المختلفة
- نقاط القوة والضعف في الموقف القانوني

#### ب. نقاط الضعف في القضية (Weak Points Analysis)
- النقاط القانونية الضعيفة في القضية أو الموقف
- الثغرات القانونية المحتملة
- الأخطاء أو المشاكل في المستندات أو الإجراءات
- المخاطر التي قد تؤثر سلباً على القضية

#### ج. نقاط القوة في القضية (Strong Points Analysis)
- النقاط القانونية القوية في القضية أو الموقف
- الحقوق الواضحة والمؤكدة
- الأدلة والوثائق الداعمة
- المزايا القانونية المتاحة

#### د. الأساس القانوني السعودي (Saudi Legal Basis)
- القوانين والأنظمة السعودية ذات الصلة
- المواد القانونية المعنية بالتفصيل (اذكر الأرقام إن أمكن)
- الأحكام القضائية المشابهة (إن وجدت)
- التفسيرات القانونية الرسمية

#### هـ. تحليل المخاطر القانونية (Legal Risk Analysis)
- المخاطر القانونية المحتملة
- احتمالية النجاح أو الفشل
- العواقب القانونية المحتملة
- التوصيات لتقليل المخاطر

#### و. الواجبات والحقوق (Obligations and Rights)
- الواجبات القانونية المترتبة على كل طرف
- الحقوق المستحقة لكل طرف
- الآليات القانونية المتاحة لحماية الحقوق

### 3. التوصيات العملية (Practical Recommendations)

#### أ. للتسوية (Settlement Recommendations)
- خيارات التسوية المناسبة
- الشروط المقترحة للتسوية
- كيفية التفاوض بشكل فعال

#### ب. للإجراءات القانونية (Legal Action Recommendations)
- الإجراءات القانونية المناسبة
- الخطوات العملية المطلوبة
- المستندات والأوراق المطلوبة
- المواعيد والجداول الزمنية

#### ج. لحماية المصالح (Protection Recommendations)
- كيفية حماية المصالح القانونية
- التدابير الوقائية المطلوبة
- إدارة المخاطر

### 4. المعلومات للعميل/المستخدم (Information for Client/User)

#### أ. شرح مبسط (Simple Explanation)
- شرح مبسط للموضوع القانوني
- ما يحتاج العميل لمعرفته
- المصطلحات القانونية شرحها بشكل مبسط

#### ب. الخطوات التالية (Next Steps)
- ما الذي يجب على العميل فعله بعد ذلك
- الوثائق المطلوبة
- المواعيد المهمة
- التحذيرات المهمة

### 5. التحليل المتقدم للمحامين (Advanced Analysis for Lawyers)

#### أ. الاستراتيجية القانونية (Legal Strategy)
- استراتيجية قانونية شاملة
- نقاط القوة والضعف التفصيلية
- السوابق القضائية ذات الصلة
- كيفية بناء الحجة القانونية

#### ب. البحث القانوني المطلوب (Required Legal Research)
- المجالات التي تحتاج لمزيد من البحث
- المصادر القانونية الإضافية
- القضايا المشابهة للرجوع إليها

#### ج. المخاطر المهنية (Professional Risks)
- المخاطر المهنية للمحامي
- الممارسات الأفضل للتعامل مع هذا النوع من القضايا
- التحذيرات الأخلاقية والمهنية

### 6. التقييم الكمي (Quantitative Assessment)
- درجة المخاطر (Risk Score): من 0 إلى 100
  - 0-25: منخفض (Low)
  - 26-50: متوسط (Medium)
  - 51-75: عالي (High)
  - 76-100: حرج (Critical)
- احتمالية النجاح (Success Probability): نسبة مئوية
- تقدير الوقت المتوقع (Estimated Timeline): بالأسابيع/الأشهر
- تقدير التكلفة (Cost Estimate): إذا كان مناسباً

### 7. المراجع القانونية (Legal References)
- قائمة بالأنظمة والقوانين السعودية ذات الصلة
- المواد القانونية المحددة
- الأحكام القضائية (إن وجدت)

---

**المتطلبات الإضافية:**
1. كن دقيقاً ومفصلاً في التحليل - لا تكتفي بالملخصات القصيرة
2. استخدم المصطلحات القانونية الصحيحة مع شرحها
3. قدم أمثلة عملية وقابلة للتطبيق
4. ركز على النظام القانوني السعودي تحديداً
5. اذكر المصادر القانونية بشكل واضح
6. قدم تحليلاً متوازناً يوضح كلا الجانبين
7. كن واقعياً في التقييمات والتوقعات
8. تحليل المستند حسب محتواه الفعلي - إذا لم يكن هناك حكم قضائي، ركز على الاستشارة والتحليل

**التنسيق المطلوب - مهم جداً:**
يجب أن يكون التحليل منقسماً إلى أقسام واضحة مع عناوين واضحة. استخدم هذا التنسيق بالضبط:

ابدأ بكتابة:
"## التحليل الكامل

### 1. ملخص تنفيذي شامل
[المحتوى هنا]

### 2. التحليل القانوني المفصل

#### أ. الوضع القانوني الحالي
[المحتوى هنا]

#### ب. نقاط الضعف في القضية
[المحتوى هنا - إذا لم ينطبق، اكتب "لا ينطبق" أو "غير متوفر"]

#### ج. نقاط القوة في القضية
[المحتوى هنا]

#### د. الأساس القانوني السعودي
[المحتوى هنا]

#### هـ. تحليل المخاطر القانونية
[المحتوى هنا]

#### و. الواجبات والحقوق
[المحتوى هنا]

### 3. التوصيات العملية
[المحتوى هنا]

### 4. المعلومات للعميل/المستخدم
[المحتوى هنا]

### 5. التحليل المتقدم للمحامين (إن أمكن)
[المحتوى هنا - إذا لم يكن هناك معلومات متقدمة، يمكن تخطيها]

### 6. التقييم الكمي
- درجة المخاطر: [رقم من 0-100]
- احتمالية النجاح: [نسبة مئوية]
- تقدير الوقت المتوقع: [مدة]

### 7. المراجع القانونية
[المحتوى هنا]"

**مهم جداً:** 
- استخدم العناوين بالضبط كما هو موضح أعلاه
- كل قسم يجب أن يكون واضحاً ومنفصلاً
- إذا كان قسم لا ينطبق (مثل الحكم القضائي إذا لم يكن موجوداً)، اذكر ذلك بوضوح
- كن مفصلاً في كل قسم - لا تكتفي بجملة واحدة

ابدأ التحليل الآن:
"""
        
        return prompt
    
    def _parse_analysis_response(self, analysis_text: str) -> Dict[str, Any]:
        """
        Parse the Gemini response into structured format.
        
        This extracts key sections from the analysis text.
        """
        # Extract risk score if mentioned
        risk_score = self._extract_risk_score(analysis_text)
        
        # Extract sections with better parsing
        sections = {
            "executive_summary": self._extract_section(analysis_text, ["1. ملخص تنفيذي", "ملخص تنفيذي", "Executive Summary", "الملخص التنفيذي"]),
            "legal_analysis": self._extract_section(analysis_text, ["2. التحليل القانوني", "التحليل القانوني المفصل", "Detailed Legal Analysis"]),
            "legal_status": self._extract_section(analysis_text, ["الوضع القانوني الحالي", "الوضع القانوني", "Legal Status"]),
            "weak_points": self._extract_section(analysis_text, ["نقاط الضعف", "نقاط الضعف في القضية", "Weak Points", "Weak Points Analysis"]),
            "strong_points": self._extract_section(analysis_text, ["نقاط القوة", "نقاط القوة في القضية", "Strong Points", "Strong Points Analysis"]),
            "legal_basis": self._extract_section(analysis_text, ["الأساس القانوني السعودي", "الأساس القانوني", "Legal Basis", "القانوني السعودي"]),
            "risk_analysis": self._extract_section(analysis_text, ["تحليل المخاطر", "تحليل المخاطر القانونية", "Risk Analysis", "المخاطر القانونية"]),
            "obligations_rights": self._extract_section(analysis_text, ["الواجبات والحقوق", "Obligations and Rights", "Obligations", "Rights"]),
            "recommendations": self._extract_section(analysis_text, ["3. التوصيات", "التوصيات العملية", "Practical Recommendations", "Recommendations"]),
            "settlement_recommendations": self._extract_section(analysis_text, ["تسوية", "للتسوية", "Settlement Recommendations", "Settlement"]),
            "legal_action_recommendations": self._extract_section(analysis_text, ["إجراءات قانونية", "للإجراءات القانونية", "Legal Action Recommendations", "Legal Action"]),
            "protection_recommendations": self._extract_section(analysis_text, ["حماية", "لحماية المصالح", "Protection Recommendations", "Protection"]),
            "client_information": self._extract_section(analysis_text, ["4. المعلومات", "المعلومات للعميل", "Information for Client", "Client", "المستخدم"]),
            "simple_explanation": self._extract_section(analysis_text, ["شرح مبسط", "Simple Explanation"]),
            "next_steps": self._extract_section(analysis_text, ["الخطوات التالية", "Next Steps"]),
            "legal_strategy": self._extract_section(analysis_text, ["5. التحليل المتقدم", "الاستراتيجية القانونية", "Legal Strategy", "Advanced Analysis"]),
            "legal_research": self._extract_section(analysis_text, ["البحث القانوني", "Legal Research"]),
            "professional_risks": self._extract_section(analysis_text, ["المخاطر المهنية", "Professional Risks"]),
            "quantitative_assessment": self._extract_section(analysis_text, ["6. التقييم الكمي", "التقييم الكمي", "Quantitative Assessment", "درجة المخاطر"]),
            "legal_references": self._extract_section(analysis_text, ["7. المراجع", "المراجع القانونية", "Legal References", "المصادر"]),
            "key_findings": self._extract_key_findings(analysis_text),
            "detailed_recommendations": self._extract_recommendations(analysis_text),
        }
        
        # Parse sections by splitting on headers - more robust method
        sections_by_header = self._parse_sections_by_headers(analysis_text)
        # Merge with keyword-based extraction (header-based takes priority)
        for key, value in sections_by_header.items():
            if value:
                sections[key] = value
        
        return {
            "risk_score": risk_score,
            "risk_label": self._get_risk_label(risk_score),
            "full_analysis": analysis_text,  # Keep full text
            "sections": sections,
            "formatted_analysis": analysis_text  # For display
        }
    
    def _extract_risk_score(self, text: str) -> int:
        """Extract risk score from analysis text."""
        import re
        # Look for patterns like "درجة المخاطر: 45" or "Risk Score: 45"
        patterns = [
            r'درجة\s*المخاطر[:\-]?\s*(\d+)',
            r'Risk\s*Score[:\-]?\s*(\d+)',
            r'المخاطر[:\-]?\s*(\d+)\s*%',
            r'Risk[:\-]?\s*(\d+)\s*%',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    return int(match.group(1))
                except ValueError:
                    continue
        
        # Default to medium risk if not found
        return 50
    
    def _extract_section(self, text: str, keywords: List[str]) -> str:
        """Extract a section from text based on keywords."""
        import re
        
        # Try multiple patterns to find sections
        for keyword in keywords:
            # Pattern 1: Look for markdown headers (### or ####)
            patterns = [
                # Markdown header pattern: ### keyword or #### keyword
                rf'(?:###|####)\s*{re.escape(keyword)}[:\-]?\s*\n+([^\n]+(?:\n(?!###|####|\d+\.)[^\n]+)*)',
                # Pattern 2: Look for numbered sections
                rf'(?:\d+\.|أ\.|ب\.|ج\.|د\.|هـ\.|و\.)\s*{re.escape(keyword)}[:\-]?\s*\n+([^\n]+(?:\n(?!\d+\.|###|####|أ\.|ب\.)[^\n]+)*)',
                # Pattern 3: Simple keyword followed by content
                rf'{re.escape(keyword)}[:\-]?\s*\n+([^\n]+(?:\n(?!###|####|\d+\.)[^\n]+)*)',
            ]
            
            for pattern in patterns:
                matches = list(re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL))
                if matches:
                    # Get the first match and extract content until next section
                    match = matches[0]
                    section_start = match.end()
                    
                    # Find the next section header or end of text
                    next_section_pattern = r'\n(?:###|####|\d+\.|أ\.|ب\.|ج\.|د\.|هـ\.|و\.)'
                    next_match = re.search(next_section_pattern, text[section_start:], re.MULTILINE)
                    
                    if next_match:
                        section_content = text[section_start:section_start + next_match.start()].strip()
                    else:
                        section_content = text[section_start:].strip()
                    
                    if section_content:
                        return section_content
        
        return ""
    
    def _extract_key_findings(self, text: str) -> List[str]:
        """Extract key findings as a list."""
        import re
        findings = []
        
        # Look for bullet points or numbered lists
        patterns = [
            r'[-•*]\s*([^\n]+)',
            r'\d+\.\s*([^\n]+)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            if matches:
                findings.extend([m.strip() for m in matches[:10]])  # Limit to 10
                break
        
        return findings[:10] if findings else []
    
    def _extract_recommendations(self, text: str) -> List[str]:
        """Extract recommendations as a list."""
        # Similar to key findings
        return self._extract_key_findings(text)
    
    def _get_risk_label(self, score: int) -> str:
        """Get risk label based on score."""
        if score <= 25:
            return "Low"
        elif score <= 50:
            return "Medium"
        elif score <= 75:
            return "High"
        else:
            return "Critical"
    
    def _parse_sections_by_headers(self, text: str) -> Dict[str, str]:
        """
        Parse analysis text by splitting on section headers.
        This is more robust than keyword matching.
        """
        import re
        sections = {}
        
        # Define section mappings (header text -> section key)
        section_mappings = {
            "1": "executive_summary",
            "ملخص تنفيذي": "executive_summary",
            "executive": "executive_summary",
            "2": "legal_analysis",
            "التحليل القانوني": "legal_analysis",
            "legal analysis": "legal_analysis",
            "الوضع القانوني": "legal_status",
            "legal status": "legal_status",
            "نقاط الضعف": "weak_points",
            "weak points": "weak_points",
            "نقاط القوة": "strong_points",
            "strong points": "strong_points",
            "الأساس القانوني": "legal_basis",
            "legal basis": "legal_basis",
            "تحليل المخاطر": "risk_analysis",
            "risk analysis": "risk_analysis",
            "الواجبات والحقوق": "obligations_rights",
            "obligations": "obligations_rights",
            "3": "recommendations",
            "التوصيات": "recommendations",
            "recommendations": "recommendations",
            "4": "client_information",
            "المعلومات": "client_information",
            "information": "client_information",
            "5": "legal_strategy",
            "التحليل المتقدم": "legal_strategy",
            "advanced": "legal_strategy",
            "6": "quantitative_assessment",
            "التقييم الكمي": "quantitative_assessment",
            "quantitative": "quantitative_assessment",
            "7": "legal_references",
            "المراجع": "legal_references",
            "references": "legal_references",
        }
        
        # Split by headers (### or #### or numbered sections like "1. " or "أ. ")
        # Pattern to match section headers
        header_pattern = r'(?:^|\n)(?:###|####)?\s*(\d+\.|أ\.|ب\.|ج\.|د\.|هـ\.|و\.)?\s*([^\n]+?)(?=\n|$)'
        
        matches = list(re.finditer(header_pattern, text, re.MULTILINE | re.IGNORECASE))
        
        for i, match in enumerate(matches):
            number_part = match.group(1) or ""
            header_text = match.group(2).strip()
            header_start = match.end()
            
            # Find the next header or end of text
            next_match = matches[i + 1] if i + 1 < len(matches) else None
            if next_match:
                section_content = text[header_start:next_match.start()].strip()
            else:
                section_content = text[header_start:].strip()
            
            # Map header to section key
            section_key = None
            header_lower = header_text.lower()
            
            # Check numbered sections first
            if number_part:
                num = re.match(r'(\d+)', number_part)
                if num and num.group(1) in section_mappings:
                    section_key = section_mappings[num.group(1)]
            
            # If not found, check header text
            if not section_key:
                for key, mapped_key in section_mappings.items():
                    if key.lower() in header_lower or header_lower in key.lower():
                        section_key = mapped_key
                        break
            
            # Store section content
            if section_key and section_content:
                # Append if section already exists (for subsections)
                if section_key in sections:
                    sections[section_key] += "\n\n" + section_content
                else:
                    sections[section_key] = section_content
        
        return sections
    
    async def _extract_text_from_docx(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from DOCX/DOC file content.
        
        Args:
            file_content: Bytes content of the DOCX/DOC file
            filename: Original filename
            
        Returns:
            Extracted text as string
        """
        if not DOCX_AVAILABLE:
            return ""
        
        try:
            # Create a temporary file to save the DOCX content
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{filename.split('.')[-1]}") as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name
            
            try:
                # Extract text using python-docx
                doc = DocxDocument(temp_path)
                text_parts = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                extracted_text = "\n\n".join(text_parts)
                
                logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX file")
                return extracted_text
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temporary file {temp_path}: {e}")
                    
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX file: {e}", exc_info=True)
            return ""

