"""
Contract Template Service for business logic.

This service handles template retrieval and contract generation.
"""

import os
import uuid
import tempfile
import json
from typing import Dict, Any, List, Optional
from pathlib import Path
from sqlalchemy.ext.asyncio import AsyncSession

from ..repositories.template_repository import TemplateRepository
from ..models.contract_template import ContractTemplate, Contract
from ..schemas.template_schemas import TemplateVariable, TemplateVariablesResponse
from ..config.enhanced_logging import get_logger

logger = get_logger(__name__)


class TemplateService:
    """Service for contract template operations."""
    
    def __init__(self, db: AsyncSession):
        """
        Initialize template service.
        
        Args:
            db: Database session
        """
        self.db = db
        self.repository = TemplateRepository(db)
        self.storage_path = Path(os.getenv("CONTRACT_STORAGE_PATH", "./storage/generated"))
        self.storage_path.mkdir(parents=True, exist_ok=True)
    
    async def preview_template(self, template_id: str) -> str:
        """
        Generate a preview of the template with empty/placeholder values.
        
        For DOCX templates, tries to generate PDF via LibreOffice/docx2pdf.
        If PDF conversion fails, generates an HTML preview instead.
        
        Args:
            template_id: Template UUID string
            
        Returns:
            Path to preview PDF or HTML file
            
        Raises:
            ValueError: If template not found
            RuntimeError: If generation fails
        """
        # Get template
        template = await self.repository.get_template_by_id(template_id)
        if not template:
            raise ValueError(f"Template {template_id} not found")
        
        if not template.is_active:
            raise ValueError(f"Template {template_id} is not active")
        
        # Create placeholder data
        variables_data = template.variables
        if isinstance(variables_data, str):
            import json
            variables_data = json.loads(variables_data)
        
        preview_data: Dict[str, Any] = {}
        for var in variables_data:
            var_name = var.get("name") if isinstance(var, dict) else getattr(var, "name", "")
            if var_name:
                # Use default value if provided, otherwise placeholder
                default = var.get("default") if isinstance(var, dict) else getattr(var, "default", None)
                var_type = var.get("type") if isinstance(var, dict) else getattr(var, "type", "text")
                
                if default:
                    preview_data[var_name] = default
                else:
                    # Use placeholder based on type
                    if var_type == "date":
                        preview_data[var_name] = "YYYY-MM-DD"
                    elif var_type == "number":
                        preview_data[var_name] = "0"
                    elif var_type == "textarea":
                        preview_data[var_name] = "---"
                    else:
                        preview_data[var_name] = "---"
        
        # Generate preview using same logic as regular generation
        try:
            logger.info(f"Generating preview for template {template_id}, format: {template.format}")
            if template.format == "docx":
                try:
                    # Try to generate PDF
                    pdf_path = await self._generate_from_docx(template, preview_data)
                    if os.path.exists(pdf_path):
                        logger.info(f"Preview PDF generated successfully: {pdf_path}")
                        return pdf_path
                    else:
                        raise FileNotFoundError(f"PDF not generated at {pdf_path}")
                except (RuntimeError, FileNotFoundError) as e:
                    # If PDF conversion fails, generate HTML preview instead
                    logger.warning(f"PDF conversion failed for preview: {str(e)}. Generating HTML preview instead.")
                    html_path = await self._generate_html_preview_from_docx(template, preview_data)
                    if os.path.exists(html_path):
                        logger.info(f"Preview HTML generated successfully: {html_path}")
                        return html_path
                    else:
                        raise FileNotFoundError(f"HTML preview not generated at {html_path}")
            elif template.format == "html":
                pdf_path = await self._generate_from_html(template, preview_data)
            else:
                raise ValueError(f"Unsupported template format: {template.format}")
            
            # Verify the file exists
            if not os.path.exists(pdf_path):
                logger.error(f"Generated preview path does not exist: {pdf_path}")
                raise FileNotFoundError(f"Generated preview file not found at: {pdf_path}")
            
            logger.info(f"Preview generated successfully: {pdf_path}")
            return pdf_path
            
        except FileNotFoundError as e:
            logger.error(f"File not found during preview generation: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Failed to generate template preview: {str(e)}", exc_info=True)
            raise RuntimeError(f"Template preview generation failed: {str(e)}")
    
    async def _generate_html_preview_from_docx(
        self,
        template: ContractTemplate,
        filled_data: Dict[str, Any]
    ) -> str:
        """
        Generate HTML preview from DOCX template (fallback when PDF conversion not available).
        
        Extracts text from DOCX, replaces variables, and creates an HTML preview.
        """
        try:
            from docx import Document
            from jinja2 import Template
            
            # Load DOCX and extract text
            if not os.path.exists(template.file_path):
                raise FileNotFoundError(f"Template file not found: {template.file_path}")
            
            doc = Document(template.file_path)
            paragraphs_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Replace variables in text
                    text = paragraph.text
                    for key, value in filled_data.items():
                        text = text.replace(f"{{{{{key}}}}}", str(value))
                    paragraphs_text.append(text)
            
            # Check if content contains Arabic
            all_text = ' '.join(paragraphs_text)
            has_arabic = any('\u0600' <= char <= '\u06FF' for char in all_text)
            direction = 'rtl' if has_arabic else 'ltr'
            text_align = 'right' if has_arabic else 'left'
            
            # Escape HTML entities for safety
            import html as html_module
            
            escaped_title = html_module.escape(str(template.title))
            escaped_description = html_module.escape(str(template.description)) if template.description else ''
            escaped_paragraphs = [html_module.escape(str(para)) for para in paragraphs_text]
            
            # Create HTML preview with Arabic support
            html_content = f"""<!DOCTYPE html>
<html dir="{direction}" lang="ar">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{escaped_title} - Preview</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+Arabic:wght@400;700&display=swap');
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: 'Noto Sans Arabic', 'Arial', 'Tahoma', sans-serif;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            line-height: 1.8;
            direction: {direction};
            text-align: {text_align};
            background: #fff;
            color: #333;
        }}
        h1 {{
            text-align: center;
            color: #333;
            font-size: 24px;
            margin-bottom: 20px;
            font-weight: bold;
        }}
        h2 {{
            color: #555;
            border-bottom: 2px solid #ddd;
            padding-bottom: 10px;
            margin-top: 30px;
            font-weight: bold;
        }}
        p {{
            margin: 10px 0;
            text-align: {text_align};
        }}
        hr {{
            margin: 20px 0;
            border: none;
            border-top: 1px solid #ddd;
        }}
    </style>
</head>
<body>
    <h1>{escaped_title}</h1>
    {f'<p style="text-align: center; color: #666;">{escaped_description}</p>' if escaped_description else ''}
    <hr>
    {''.join(f'<p>{para}</p>' for para in escaped_paragraphs)}
</body>
</html>"""
            
            # Save HTML preview
            html_filename = f"preview_{uuid.uuid4()}.html"
            html_path = self.storage_path / html_filename
            
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return str(html_path)
            
        except Exception as e:
            logger.error(f"Failed to generate HTML preview: {str(e)}")
            raise RuntimeError(f"HTML preview generation failed: {str(e)}")
    
    async def get_template_variables(self, template_id: str) -> TemplateVariablesResponse:
        """
        Get template variables schema for form generation.
        
        Args:
            template_id: Template UUID string
            
        Returns:
            TemplateVariablesResponse with template info and variables
            
        Raises:
            ValueError: If template not found
        """
        template = await self.repository.get_template_by_id(template_id)
        if not template:
            raise ValueError(f"Template {template_id} not found")
        
        if not template.is_active:
            raise ValueError(f"Template {template_id} is not active")
        
        # Parse variables from JSON
        variables_data = template.variables
        if isinstance(variables_data, str):
            variables_data = json.loads(variables_data)
        
        variables = [
            TemplateVariable(**var) if isinstance(var, dict) else var
            for var in variables_data
        ]
        
        return TemplateVariablesResponse(
            id=template.id,
            title=template.title,
            description=template.description,
            variables=variables
        )
    
    async def generate_contract(
        self,
        template_id: str,
        owner_id: int,
        filled_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Generate a contract PDF from template and user data.
        
        Args:
            template_id: Template UUID string
            owner_id: User ID who owns the contract
            filled_data: User-provided form data
            
        Returns:
            Dict with contract_id and pdf_url
            
        Raises:
            ValueError: If template not found or validation fails
            RuntimeError: If generation fails
        """
        # Get template
        template = await self.repository.get_template_by_id(template_id)
        if not template:
            raise ValueError(f"Template {template_id} not found")
        
        if not template.is_active:
            raise ValueError(f"Template {template_id} is not active")
        
        # Validate required fields
        variables_data = template.variables
        if isinstance(variables_data, str):
            variables_data = json.loads(variables_data)
        
        required_fields = [
            var.get("name") for var in variables_data
            if var.get("required", False)
        ]
        
        missing_fields = [
            field for field in required_fields
            if field not in filled_data or not filled_data.get(field)
        ]
        
        if missing_fields:
            raise ValueError(f"Missing required fields: {', '.join(missing_fields)}")
        
        # Generate contract based on format
        try:
            if template.format == "docx":
                pdf_path = await self._generate_from_docx(template, filled_data)
            elif template.format == "html":
                pdf_path = await self._generate_from_html(template, filled_data)
            else:
                raise ValueError(f"Unsupported template format: {template.format}")
            
            # Ensure we store absolute path for reliable file access
            pdf_path_abs = os.path.abspath(str(pdf_path))
            
            # Verify file exists before storing
            if not os.path.exists(pdf_path_abs):
                logger.error(f"Generated file does not exist at: {pdf_path_abs}")
                raise RuntimeError(f"Generated contract file not found at: {pdf_path_abs}")
            
            # Determine file type for the response message
            file_extension = Path(pdf_path_abs).suffix.lower()
            is_pdf = file_extension == '.pdf'
            is_docx = file_extension == '.docx'
            
            logger.info(f"Storing contract with file path: {pdf_path_abs} (exists: {os.path.exists(pdf_path_abs)})")
            
            # Create contract record with absolute path
            contract = await self.repository.create_contract(
                template_id=template_id,
                owner_id=owner_id,
                filled_data=filled_data,
                pdf_path=pdf_path_abs,  # Store absolute path
                status="generated"
            )
            
            await self.db.commit()
            
            # Generate URL (adjust based on your deployment)
            base_url = os.getenv("BACKEND_URL", "http://localhost:8000")
            download_url = f"{base_url}/api/v1/templates/contracts/{contract.id}/download"
            
            result = {
                "contract_id": contract.id,
                "pdf_url": download_url,  # Keep pdf_url for backwards compatibility
                "download_url": download_url,
                "success": True
            }
            
            # Add file type information
            if is_docx:
                result["file_type"] = "docx"
                result["message"] = "Contract generated successfully (DOCX format - PDF conversion not available)"
                logger.warning(f"Contract {contract.id} generated as DOCX instead of PDF. LibreOffice/Word required for PDF conversion.")
            elif is_pdf:
                result["file_type"] = "pdf"
            
            return result
            
        except Exception as e:
            await self.db.rollback()
            logger.error(f"Failed to generate contract: {str(e)}")
            raise RuntimeError(f"Contract generation failed: {str(e)}")
    
    async def _generate_from_docx(
        self,
        template: ContractTemplate,
        filled_data: Dict[str, Any]
    ) -> str:
        """
        Generate PDF from DOCX template using docxtpl.
        
        Args:
            template: ContractTemplate model
            filled_data: User-provided form data
            
        Returns:
            Path to generated PDF
            
        Note: Requires docxtpl library for advanced template rendering.
        Install with: pip install docxtpl
        Falls back to python-docx for basic text replacement if not available.
        """
        try:
            from docxtpl import DocxTemplate
            
            # Check if file exists
            if not os.path.exists(template.file_path):
                raise FileNotFoundError(f"Template file not found at: {template.file_path}")
            
            # Process Arabic text in filled_data if needed
            processed_data = {}
            for key, value in filled_data.items():
                if value and isinstance(value, str):
                    # Check if value contains Arabic
                    has_arabic = any('\u0600' <= char <= '\u06FF' for char in str(value))
                    if has_arabic:
                        # Apply Arabic text processing if available
                        try:
                            import arabic_reshaper
                            from bidi.algorithm import get_display
                            reshaped = arabic_reshaper.reshape(str(value))
                            processed_data[key] = get_display(reshaped)
                        except ImportError:
                            processed_data[key] = value
                    else:
                        processed_data[key] = value
                else:
                    processed_data[key] = value
            
            # Load DOCX template
            doc = DocxTemplate(template.file_path)
            
            # Render with processed data
            doc.render(processed_data)
            
            # Save to temp DOCX
            temp_docx = tempfile.NamedTemporaryFile(
                suffix=".docx",
                delete=False,
                dir=self.storage_path
            )
            doc.save(temp_docx.name)
            temp_docx.close()
            
            # Convert DOCX to PDF using LibreOffice (or alternative)
            try:
                pdf_path = await self._convert_docx_to_pdf(temp_docx.name)
                # Clean up temp DOCX if PDF conversion succeeded
                os.unlink(temp_docx.name)
                return pdf_path
            except RuntimeError as e:
                # PDF conversion failed - return DOCX as fallback
                logger.warning(f"PDF conversion failed, returning DOCX file instead: {str(e)}")
                # Keep the DOCX file and return its path
                # Rename to a proper location instead of temp
                docx_filename = f"{uuid.uuid4()}.docx"
                docx_path = self.storage_path / docx_filename
                import shutil
                shutil.move(temp_docx.name, str(docx_path))
                logger.info(f"Generated DOCX file at: {docx_path}")
                return str(docx_path)
            
        except ImportError:
            # Fallback: use python-docx if docxtpl not available
            logger.warning("docxtpl not available, using basic python-docx")
            from docx import Document
            
            doc = Document(template.file_path)
            
            # Simple text replacement (for basic templates)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    for key, value in filled_data.items():
                        if f"{{{{{key}}}}}" in run.text:
                            run.text = run.text.replace(f"{{{{{key}}}}}", str(value))
            
            # Save and convert
            temp_docx = tempfile.NamedTemporaryFile(
                suffix=".docx",
                delete=False,
                dir=self.storage_path
            )
            doc.save(temp_docx.name)
            temp_docx.close()
            
            try:
                pdf_path = await self._convert_docx_to_pdf(temp_docx.name)
                os.unlink(temp_docx.name)
                return pdf_path
            except RuntimeError as e:
                # PDF conversion failed - return DOCX as fallback
                logger.warning(f"PDF conversion failed, returning DOCX file instead: {str(e)}")
                docx_filename = f"{uuid.uuid4()}.docx"
                docx_path = self.storage_path / docx_filename
                import shutil
                shutil.move(temp_docx.name, str(docx_path))
                logger.info(f"Generated DOCX file at: {docx_path}")
                return str(docx_path)
    
    async def _generate_from_html(
        self,
        template: ContractTemplate,
        filled_data: Dict[str, Any]
    ) -> str:
        """
        Generate PDF from HTML template using Jinja2 + WeasyPrint.
        
        Args:
            template: ContractTemplate model
            filled_data: User-provided form data
            
        Returns:
            Path to generated PDF
            
        Note: Requires weasyprint library for PDF generation from HTML.
        Install with: pip install weasyprint
        Falls back to HTML file generation if not available.
        """
        try:
            from jinja2 import Environment, FileSystemLoader, Template
            from weasyprint import HTML
            
            # Load HTML template
            with open(template.file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Render with Jinja2
            env = Environment(loader=FileSystemLoader(os.path.dirname(template.file_path)))
            jinja_template = env.from_string(html_content)
            rendered_html = jinja_template.render(**filled_data)
            
            # Generate PDF
            pdf_filename = f"{uuid.uuid4()}.pdf"
            pdf_path = self.storage_path / pdf_filename
            
            HTML(string=rendered_html).write_pdf(str(pdf_path))
            
            return str(pdf_path)
            
        except ImportError:
            # Fallback: Generate HTML file (user can print to PDF)
            logger.warning("WeasyPrint not available, generating HTML file")
            
            from jinja2 import Environment, Template
            
            with open(template.file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            template_obj = Template(html_content)
            rendered_html = template_obj.render(**filled_data)
            
            html_filename = f"{uuid.uuid4()}.html"
            html_path = self.storage_path / html_filename
            
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(rendered_html)
            
            # Return HTML path (frontend can handle conversion)
            return str(html_path)
    
    async def _convert_docx_to_pdf(self, docx_path: str) -> str:
        """
        Convert DOCX file to PDF.
        
        Uses LibreOffice command line if available, otherwise uses docx2pdf or returns DOCX path.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Path to PDF file
        """
        import subprocess
        import platform
        import asyncio
        
        # Ensure docx_path is a string and normalize path separators
        docx_path = str(docx_path).replace('/', os.sep)
        
        # Generate PDF path in the same directory as DOCX
        docx_path_obj = Path(docx_path)
        pdf_path = str(docx_path_obj.with_suffix('.pdf'))
        
        # Determine LibreOffice command based on OS
        if platform.system() == "Windows":
            libreoffice_cmds = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                "soffice",
            ]
        else:
            libreoffice_cmds = ["libreoffice", "soffice"]
        
        # Try LibreOffice conversion
        for cmd in libreoffice_cmds:
            try:
                result = subprocess.run(
                    [
                        cmd,
                        "--headless",
                        "--convert-to", "pdf",
                        "--outdir", str(self.storage_path),
                        docx_path
                    ],
                    check=True,
                    timeout=60,
                    capture_output=True,
                    text=True
                )
                
                # LibreOffice outputs PDF to outdir with same base name
                expected_pdf = Path(self.storage_path) / Path(docx_path).stem
                expected_pdf = expected_pdf.with_suffix(".pdf")
                
                if os.path.exists(str(expected_pdf)):
                    logger.info(f"Successfully converted DOCX to PDF: {expected_pdf}")
                    return str(expected_pdf)
                elif os.path.exists(pdf_path):
                    logger.info(f"Successfully converted DOCX to PDF: {pdf_path}")
                    return pdf_path
                    
            except FileNotFoundError:
                continue  # Try next command
            except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
                logger.warning(f"LibreOffice conversion failed with {cmd}: {str(e)}")
                if hasattr(e, 'stderr') and e.stderr:
                    logger.warning(f"LibreOffice stderr: {e.stderr}")
                continue  # Try next command
        
        # Try docx2pdf as fallback (Python library)
        # Run in executor since convert() is blocking
        try:
            from docx2pdf import convert
            logger.info(f"Attempting docx2pdf conversion: {docx_path} -> {pdf_path}")
            
            # Ensure paths are absolute and normalized
            docx_abs = os.path.abspath(docx_path)
            pdf_abs = os.path.abspath(pdf_path)
            
            logger.debug(f"Absolute paths - DOCX: {docx_abs}, PDF: {pdf_abs}")
            
            # Run convert in executor to avoid blocking the event loop
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, convert, docx_abs, pdf_abs)
            
            # Check for PDF in multiple possible locations
            possible_pdf_paths = [
                pdf_abs,
                pdf_path,
                str(Path(docx_abs).with_suffix('.pdf')),
                os.path.join(os.path.dirname(docx_abs), Path(docx_abs).stem + '.pdf')
            ]
            
            pdf_found = None
            for possible_path in possible_pdf_paths:
                if os.path.exists(possible_path):
                    pdf_found = possible_path
                    break
            
            if pdf_found:
                logger.info(f"Successfully converted DOCX to PDF using docx2pdf: {pdf_found}")
                return pdf_found
            else:
                logger.error(f"docx2pdf convert() completed but PDF file not found")
                logger.error(f"Checked paths: {possible_pdf_paths}")
                raise RuntimeError(f"PDF file not created. Expected at: {pdf_abs}")
        except ImportError:
            logger.warning("docx2pdf not available. Install with: pip install docx2pdf")
        except Exception as e:
            error_msg = str(e)
            error_type = type(e).__name__
            logger.error(f"docx2pdf conversion failed [Type: {error_type}]: {error_msg}")
            # Log full exception details for debugging
            import traceback
            logger.debug(f"docx2pdf full traceback:\n{traceback.format_exc()}")
            
            # On Windows, docx2pdf requires LibreOffice or MS Word
            if platform.system() == "Windows":
                if "LibreOffice" in error_msg or "soffice" in error_msg.lower():
                    logger.warning("docx2pdf requires LibreOffice on Windows. Install from https://www.libreoffice.org/")
                elif "Word" in error_msg or "COM" in error_msg or "comtypes" in error_msg.lower():
                    logger.warning("docx2pdf COM conversion failed. Ensure Word is installed and accessible via COM")
                    logger.warning("Try: 1) Close any open Word instances, 2) Restart the Python server")
                elif "Permission" in error_msg or "access" in error_msg.lower():
                    logger.warning("Permission error accessing Word. Try running as administrator or check COM permissions")
        
        # If all conversion methods fail, raise an error with helpful message
        logger.error("All PDF conversion methods failed.")
        
        # Provide OS-specific installation instructions
        if platform.system() == "Windows":
            error_message = (
                "PDF conversion failed. On Windows, you need to install either:\n"
                "1. LibreOffice (recommended): Download from https://www.libreoffice.org/\n"
                "2. Microsoft Word (for COM-based conversion)\n"
                "Then install the Python package: pip install docx2pdf"
            )
        else:
            error_message = (
                "PDF conversion failed. Please install LibreOffice:\n"
                "Ubuntu/Debian: sudo apt-get install libreoffice\n"
                "macOS: brew install libreoffice\n"
                "Or install docx2pdf Python package: pip install docx2pdf"
            )
        
        raise RuntimeError(error_message)

