"""
Seed script to create test contract templates in the database.
This script creates sample templates that match the frontend mock data.
"""

import asyncio
import uuid
import json
import logging
from pathlib import Path
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.database import AsyncSessionLocal
from app.models.contract_template import ContractTemplate

logger = logging.getLogger(__name__)


async def create_template(
    db: AsyncSession,
    title: str,
    description: str,
    category: str,
    variables: list,
    file_path: str = None,
    format: str = "docx"
) -> ContractTemplate:
    """Create a template in the database."""
    # Generate a simple template file if it doesn't exist
    if not file_path or not Path(file_path).exists():
        # Create a simple DOCX or HTML template file
        storage_dir = Path("./storage/templates")
        storage_dir.mkdir(parents=True, exist_ok=True)
        
        if format == "docx":
            # Create a proper DOCX file using python-docx
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                file_path = str(storage_dir / f"{title.lower().replace(' ', '_')}.docx")
                doc = Document()
                
                # Add title
                title_para = doc.add_heading(title, 0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add description
                if description:
                    desc_para = doc.add_paragraph(description)
                    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()  # Empty line
                
                # Add template content with variables in Arabic
                doc.add_heading('شروط العقد', level=1)
                doc.add_paragraph("تم إبرام هذه الاتفاقية في تاريخ {{{{ start_date }}}} بين الطرف الأول: {{{{ partyA_name }}}} والطرف الثاني: {{{{ partyB_name }}}}.")
                doc.add_paragraph()
                
                # Add sections based on template type
                for var in variables:
                    if var['type'] != 'date' and 'name' in var and 'label' in var:
                        # Map English labels to Arabic
                        arabic_labels = {
                            'Employer Name': 'اسم صاحب العمل',
                            'Employee Name': 'اسم الموظف',
                            'Job Position': 'المسمى الوظيفي',
                            'Salary Amount': 'الراتب',
                            'Disclosing Party': 'الطرف الأول',
                            'Receiving Party': 'الطرف الثاني',
                            'Service Provider': 'مقدم الخدمة',
                            'Client Name': 'اسم العميل',
                            'Seller Name': 'اسم البائع',
                            'Buyer Name': 'اسم المشتري',
                            'Contract Amount': 'قيمة العقد',
                            'Purchase Amount': 'قيمة الشراء',
                            'Duration (Months)': 'المدة (بالأشهر)',
                            'Description of Confidential Information': 'وصف المعلومات السرية',
                            'Service Description': 'وصف الخدمة',
                            'Uptime Percentage (%)': 'نسبة الوقت المتاح (%)',
                            'Item Description': 'وصف البضاعة',
                        }
                        label = arabic_labels.get(var['label'], var['label'])
                        doc.add_paragraph(f"{label}: {{{{ {var['name']} }}}}")
                
                doc.add_paragraph()
                doc.add_paragraph("يتفق الطرفان على الشروط والأحكام الواردة أعلاه.")
                
                # Save the document
                doc.save(file_path)
            except ImportError:
                # Fallback: Create HTML template if python-docx not available
                logger.warning("python-docx not available, creating HTML template instead")
                file_path = str(storage_dir / f"{title.lower().replace(' ', '_')}.html")
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(f"<html><body><h1>{title}</h1>")
                    f.write(f"<p>{description}</p>")
                    for var in variables:
                        f.write(f"<p>{{{{ {var['name']} }}}}</p>")
                    f.write("</body></html>")
        else:
            file_path = str(storage_dir / f"{title.lower().replace(' ', '_')}.html")
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"<html><body><h1>{title}</h1>")
                f.write(f"<p>{description}</p>")
                for var in variables:
                    f.write(f"<p>{{{{ {var['name']} }}}}</p>")
                f.write("</body></html>")
    
    # Generate UUID that matches frontend mock data IDs (optional mapping)
    template_id = str(uuid.uuid4())
    
    template = ContractTemplate(
        id=template_id,
        title=title,
        description=description,
        category=category,
        file_path=file_path,
        format=format,
        variables=variables,
        is_active=True,
        is_premium=False
    )
    
    db.add(template)
    await db.flush()
    await db.refresh(template)
    
    return template


async def seed_templates():
    """Seed the database with test templates."""
    print("=" * 60)
    print("  Contract Templates Seeding Script")
    print("=" * 60)
    print()
    
    async with AsyncSessionLocal() as db:
        try:
            # Template 1: Employment Agreement
            template1 = await create_template(
                db=db,
                title="Employment Agreement Template",
                description="Standard employment agreement with customizable terms",
                category="Employment",
                format="docx",
                variables=[
                    {
                        "name": "partyA_name",
                        "label": "Employer Name",
                        "type": "text",
                        "required": True,
                        "placeholder": "Enter company name"
                    },
                    {
                        "name": "partyB_name",
                        "label": "Employee Name",
                        "type": "text",
                        "required": True,
                        "placeholder": "Enter employee name"
                    },
                    {
                        "name": "start_date",
                        "label": "Start Date",
                        "type": "date",
                        "required": True
                    },
                    {
                        "name": "salary",
                        "label": "Salary Amount",
                        "type": "number",
                        "required": False,
                        "placeholder": "Enter salary amount"
                    },
                    {
                        "name": "position",
                        "label": "Job Position",
                        "type": "text",
                        "required": True,
                        "placeholder": "Enter job title"
                    }
                ]
            )
            print(f"✅ Created template: {template1.title} (ID: {template1.id})")
            
            # Template 2: Non-Disclosure Agreement
            template2 = await create_template(
                db=db,
                title="Non-Disclosure Agreement (NDA)",
                description="Comprehensive NDA template for business partnerships",
                category="Confidentiality",
                format="docx",
                variables=[
                    {
                        "name": "partyA_name",
                        "label": "Party A Name",
                        "type": "text",
                        "required": True,
                        "placeholder": "Enter party A name"
                    },
                    {
                        "name": "partyB_name",
                        "label": "Party B Name",
                        "type": "text",
                        "required": True,
                        "placeholder": "Enter party B name"
                    },
                    {
                        "name": "effective_date",
                        "label": "Effective Date",
                        "type": "date",
                        "required": True
                    },
                    {
                        "name": "duration_months",
                        "label": "Duration (Months)",
                        "type": "number",
                        "required": True,
                        "placeholder": "Enter duration in months"
                    },
                    {
                        "name": "purpose",
                        "label": "Purpose",
                        "type": "text",
                        "required": True,
                        "placeholder": "Enter the purpose of disclosure"
                    },
                    {
                        "name": "post_duration_months",
                        "label": "Post-Term Duration (Months)",
                        "type": "number",
                        "required": True,
                        "placeholder": "Enter post-agreement confidentiality duration"
                    }
                ]
            )
            print(f"✅ Created template: {template2.title} (ID: {template2.id})")
            
            # Template 3: Service Level Agreement
            template3 = await create_template(
                db=db,
                title="Service Level Agreement",
                description="SLA template for IT services and support agreements",
                category="Service",
                format="docx",
                variables=[
                    {
                        "name": "effective_date",
                        "label": "Effective Date",
                        "type": "date",
                        "required": True
                    },
                    {
                        "name": "service_provider",
                        "label": "Service Provider",
                        "type": "text",
                        "required": True,
                        "placeholder": "Enter service provider name"
                    },
                    {
                        "name": "client_name",
                        "label": "Client Name",
                        "type": "text",
                        "required": True,
                        "placeholder": "Enter client name"
                    },
                    {
                        "name": "service_description",
                        "label": "Service Description",
                        "type": "textarea",
                        "required": True,
                        "placeholder": "Describe the service"
                    },
                    {
                        "name": "service_start_date",
                        "label": "Service Start Date",
                        "type": "date",
                        "required": True
                    },
                    {
                        "name": "service_end_date",
                        "label": "Service End Date",
                        "type": "date",
                        "required": True
                    },
                    {
                        "name": "service_amount",
                        "label": "Service Amount (SAR)",
                        "type": "number",
                        "required": True,
                        "placeholder": "Enter service amount in SAR"
                    },
                    {
                        "name": "notice_period",
                        "label": "Notice Period (Days)",
                        "type": "number",
                        "required": True,
                        "placeholder": "Enter notice period in days"
                    },
                    {
                        "name": "uptime_percentage",
                        "label": "Uptime Percentage (%)",
                        "type": "number",
                        "required": False,
                        "placeholder": "e.g., 99.9"
                    },
                    {
                        "name": "response_time",
                        "label": "Response Time (Minutes)",
                        "type": "number",
                        "required": False,
                        "placeholder": "Enter response time in minutes"
                    },
                    {
                        "name": "resolution_time",
                        "label": "Resolution Time (Hours)",
                        "type": "number",
                        "required": False,
                        "placeholder": "Enter resolution time in hours"
                    },
                    {
                        "name": "duration_months",
                        "label": "Duration (Months)",
                        "type": "number",
                        "required": True,
                        "placeholder": "Enter duration in months"
                    }
                ]
            )
            print(f"✅ Created template: {template3.title} (ID: {template3.id})")
            
            # Template 4: Purchase Agreement
            template4 = await create_template(
                db=db,
                title="Purchase Agreement Template",
                description="Standard purchase agreement for goods and services",
                category="Commercial",
                format="docx",
                variables=[
                    {
                        "name": "seller_name",
                        "label": "Seller Name",
                        "type": "text",
                        "required": True,
                        "placeholder": "Enter seller name"
                    },
                    {
                        "name": "buyer_name",
                        "label": "Buyer Name",
                        "type": "text",
                        "required": True,
                        "placeholder": "Enter buyer name"
                    },
                    {
                        "name": "purchase_date",
                        "label": "Purchase Date",
                        "type": "date",
                        "required": True
                    },
                    {
                        "name": "item_description",
                        "label": "Item Description",
                        "type": "textarea",
                        "required": True,
                        "placeholder": "Describe the items being purchased"
                    },
                    {
                        "name": "purchase_amount",
                        "label": "Purchase Amount",
                        "type": "number",
                        "required": True,
                        "placeholder": "Enter purchase amount"
                    },
                    {
                        "name": "payment_terms",
                        "label": "Payment Terms (Days)",
                        "type": "number",
                        "required": True,
                        "placeholder": "Enter payment terms in days"
                    },
                    {
                        "name": "delivery_location",
                        "label": "Delivery Location",
                        "type": "text",
                        "required": True,
                        "placeholder": "Enter delivery location"
                    },
                    {
                        "name": "delivery_date",
                        "label": "Delivery Date",
                        "type": "date",
                        "required": True
                    },
                    {
                        "name": "warranty_period",
                        "label": "Warranty Period (Days)",
                        "type": "number",
                        "required": True,
                        "placeholder": "Enter warranty period in days"
                    }
                ]
            )
            print(f"✅ Created template: {template4.title} (ID: {template4.id})")
            
            await db.commit()
            
            print("\n" + "=" * 60)
            print("✨ Templates seeded successfully!")
            print("=" * 60)
            print("\n📋 Created Templates:")
            print(f"  1. {template1.title} - ID: {template1.id}")
            print(f"  2. {template2.title} - ID: {template2.id}")
            print(f"  3. {template3.title} - ID: {template3.id}")
            print(f"  4. {template4.title} - ID: {template4.id}")
            print("\n💡 Note: Use these template IDs in the frontend to test the contract generator.")
            print("💡 Template files are created in ./storage/templates/ directory.")
            
            return [template1, template2, template3, template4]
            
        except Exception as e:
            print(f"\n❌ Error: {str(e)}")
            import traceback
            traceback.print_exc()
            await db.rollback()
            raise


if __name__ == "__main__":
    asyncio.run(seed_templates())

