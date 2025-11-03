"""
Script to fix existing template files - convert .txt to .docx or create new DOCX files.
"""

import asyncio
import logging
from pathlib import Path
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.database import AsyncSessionLocal
from app.models.contract_template import ContractTemplate
from sqlalchemy import select

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


async def fix_template_files():
    """Fix template files by creating proper DOCX files."""
    print("=" * 60)
    print("  Fix Template Files Script")
    print("=" * 60)
    print()
    
    async with AsyncSessionLocal() as db:
        try:
            # Get all templates
            result = await db.execute(select(ContractTemplate))
            templates = result.scalars().all()
            
            if not templates:
                print("No templates found in database.")
                return
            
            print(f"Found {len(templates)} templates to fix.\n")
            
            storage_dir = Path("./storage/templates")
            storage_dir.mkdir(parents=True, exist_ok=True)
            
            for template in templates:
                try:
                    # Check if current file exists and is valid
                    current_path = Path(template.file_path)
                    
                    # Always regenerate to ensure Arabic content (force update)
                    # If file doesn't exist or is .txt, create new DOCX
                    if not current_path.exists() or current_path.suffix == '.txt' or True:  # Force regeneration
                        print(f"Fixing template: {template.title}")
                        
                        # Create proper DOCX file
                        try:
                            from docx import Document
                            from docx.enum.text import WD_ALIGN_PARAGRAPH
                            
                            new_file_path = storage_dir / f"{template.title.lower().replace(' ', '_').replace('/', '_')}.docx"
                            
                            doc = Document()
                            
                            # Add title
                            title_para = doc.add_heading(template.title, 0)
                            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Add description
                            if template.description:
                                desc_para = doc.add_paragraph(template.description)
                                desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                doc.add_paragraph()
                            
                            # Parse variables
                            variables_data = template.variables
                            if isinstance(variables_data, str):
                                import json
                                variables_data = json.loads(variables_data)
                            
                            # Add contract content with variables in Arabic
                            doc.add_heading('شروط العقد', level=1)
                            doc.add_paragraph("تم إبرام هذه الاتفاقية بين الطرفين المحددين أدناه.")
                            doc.add_paragraph()
                            
                            # Arabic labels mapping
                            arabic_labels = {
                                'Employer Name': 'اسم صاحب العمل',
                                'Employee Name': 'اسم الموظف',
                                'Job Position': 'المسمى الوظيفي',
                                'Salary Amount': 'الراتب',
                                'Start Date': 'تاريخ البدء',
                                'Disclosing Party': 'الطرف المكشف',
                                'Receiving Party': 'الطرف المستقبل',
                                'Effective Date': 'تاريخ سريان العقد',
                                'Duration (Months)': 'المدة (بالأشهر)',
                                'Service Provider': 'مقدم الخدمة',
                                'Client Name': 'اسم العميل',
                                'Service Start Date': 'تاريخ بدء الخدمة',
                                'Service Description': 'وصف الخدمة',
                                'Uptime Percentage (%)': 'نسبة الوقت المتاح (%)',
                                'Seller Name': 'اسم البائع',
                                'Buyer Name': 'اسم المشتري',
                                'Purchase Date': 'تاريخ الشراء',
                                'Item Description': 'وصف البضاعة',
                                'Purchase Amount': 'قيمة الشراء',
                                'Description of Confidential Information': 'وصف المعلومات السرية',
                            }
                            
                            # Add full Arabic contract content based on template type
                            if template.category == 'Employment':
                                # Employment contract will use generic format
                                doc.add_paragraph("عقد عمل")
                                doc.add_paragraph()
                                doc.add_paragraph(f"صاحب العمل: {{{{ partyA_name }}}}")
                                doc.add_paragraph(f"الموظف: {{{{ partyB_name }}}}")
                                doc.add_paragraph(f"تاريخ البدء: {{{{ start_date }}}}")
                                doc.add_paragraph(f"المسمى الوظيفي: {{{{ position }}}}")
                                doc.add_paragraph(f"الراتب: {{{{ salary }}}}")
                            elif template.category == 'Confidentiality':
                                # Full NDA template
                                doc.add_paragraph("اتفاقية عدم إفصاح")
                                doc.add_paragraph()
                                doc.add_paragraph(f"تم إبرام هذه الاتفاقية بتاريخ {{{{ effective_date }}}} بين كل من:")
                                doc.add_paragraph()
                                doc.add_paragraph(f"1. {{{{ partyA_name }}}} (\"الطرف الأول\").")
                                doc.add_paragraph(f"2. {{{{ partyB_name }}}} (\"الطرف الثاني\").")
                                doc.add_paragraph()
                                doc.add_paragraph("وحيث أن الطرف الأول يملك معلومات سرية يرغب في حمايتها، فقد تم الاتفاق على ما يلي:")
                                doc.add_paragraph()
                                doc.add_paragraph("المادة (1): التعريف")
                                doc.add_paragraph("المعلومات السرية تشمل كافة البيانات والمستندات والأفكار التي يكشف عنها الطرف الأول للطرف الثاني، سواء كانت مكتوبة أو شفهية.")
                                doc.add_paragraph()
                                doc.add_paragraph("المادة (2): الالتزام بالسرية")
                                doc.add_paragraph(f"يتعهد الطرف الثاني بالحفاظ على سرية المعلومات وعدم استخدامها لأي غرض آخر غير الغرض المحدد وهو {{{{ purpose }}}}.")
                                doc.add_paragraph()
                                doc.add_paragraph("المادة (3): المدة")
                                doc.add_paragraph(f"تسري هذه الاتفاقية لمدة {{{{ duration_months }}}} أشهر من تاريخ توقيعها، ويستمر الالتزام بالسرية لمدة {{{{ post_duration_months }}}} أشهر بعد انتهائها.")
                                doc.add_paragraph()
                                doc.add_paragraph("المادة (4): الاستثناءات")
                                doc.add_paragraph("لا تنطبق السرية على المعلومات التي أصبحت متاحة للعامة دون خرق لهذه الاتفاقية.")
                                doc.add_paragraph()
                                doc.add_paragraph("المادة (5): القانون المطبق")
                                doc.add_paragraph("تخضع هذه الاتفاقية لأنظمة المملكة العربية السعودية.")
                                doc.add_paragraph()
                                doc.add_paragraph("التوقيع:")
                                doc.add_paragraph("الطرف الأول: _____________________")
                                doc.add_paragraph("الطرف الثاني: ___________________")
                            elif template.category == 'Service':
                                # Check if it's SLA or Service Agreement based on title
                                if 'SLA' in template.title or 'Service Level' in template.title:
                                    # SLA Template
                                    doc.add_paragraph("اتفاقية مستوى الخدمة (SLA)")
                                    doc.add_paragraph()
                                    doc.add_paragraph(f"تم بتاريخ {{{{ effective_date }}}} بين:")
                                    doc.add_paragraph()
                                    doc.add_paragraph(f"1. {{{{ service_provider }}}} (مقدم الخدمة)")
                                    doc.add_paragraph(f"2. {{{{ client_name }}}} (العميل)")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (1): الغرض")
                                    doc.add_paragraph("تحدد هذه الاتفاقية معايير الأداء والالتزامات التشغيلية لتقديم الخدمة.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (2): تعريف الخدمة")
                                    doc.add_paragraph(f"الخدمة المقدمة هي: {{{{ service_description }}}}.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (3): مستوى الأداء")
                                    doc.add_paragraph(f"- نسبة الوقت المتاح: {{{{ uptime_percentage }}}}%.")
                                    doc.add_paragraph(f"- زمن الاستجابة: {{{{ response_time }}}} دقيقة.")
                                    doc.add_paragraph(f"- زمن الحل للمشكلات الحرجة: {{{{ resolution_time }}}} ساعة.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (4): التقارير والمراجعة")
                                    doc.add_paragraph("يلتزم مقدم الخدمة بتزويد العميل بتقرير أداء شهري موضح فيه مؤشرات الخدمة.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (5): الإخلال بالاتفاق")
                                    doc.add_paragraph("في حال عدم الالتزام بالمستوى المتفق عليه، يحق للعميل المطالبة بالتعويض أو الخصم المالي حسب الاتفاق.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (6): مدة الاتفاق")
                                    doc.add_paragraph(f"تسري الاتفاقية لمدة {{{{ duration_months }}}} شهراً من تاريخ التوقيع.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (7): القانون المطبق")
                                    doc.add_paragraph("تخضع هذه الاتفاقية لأنظمة المملكة العربية السعودية.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("التوقيع:")
                                    doc.add_paragraph("مقدم الخدمة: _____________________")
                                    doc.add_paragraph("العميل: _________________________")
                                else:
                                    # General Service Agreement template
                                    doc.add_paragraph("اتفاقية تقديم خدمات")
                                    doc.add_paragraph()
                                    doc.add_paragraph(f"تم بعون الله بتاريخ {{{{ effective_date }}}} بين:")
                                    doc.add_paragraph()
                                    doc.add_paragraph(f"1. {{{{ service_provider }}}}، ويشار إليه بـ \"مقدم الخدمة\".")
                                    doc.add_paragraph(f"2. {{{{ client_name }}}}، ويشار إليه بـ \"العميل\".")
                                    doc.add_paragraph()
                                    doc.add_paragraph("تم الاتفاق على ما يلي:")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (1): نطاق الخدمة")
                                    doc.add_paragraph(f"يقدم مقدم الخدمة للعميل الخدمات التالية: {{{{ service_description }}}}.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (2): مدة العقد")
                                    doc.add_paragraph(f"تبدأ الخدمة من تاريخ {{{{ service_start_date }}}} وتنتهي بتاريخ {{{{ service_end_date }}}}.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (3): المقابل المالي")
                                    doc.add_paragraph(f"يتقاضى مقدم الخدمة مبلغ {{{{ service_amount }}}} ريال سعودي نظير تقديم الخدمات المتفق عليها.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (4): التزامات مقدم الخدمة")
                                    doc.add_paragraph("- الالتزام بالمواعيد المحددة.")
                                    doc.add_paragraph("- تنفيذ الخدمة وفق المعايير المهنية.")
                                    doc.add_paragraph("- إخطار العميل فور حدوث أي عائق.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (5): التزامات العميل")
                                    doc.add_paragraph("- تسليم كافة البيانات المطلوبة.")
                                    doc.add_paragraph("- سداد المستحقات المالية في المواعيد المحددة.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (6): الإنهاء")
                                    doc.add_paragraph(f"يجوز لأي من الطرفين إنهاء العقد بإشعار كتابي قبل {{{{ notice_period }}}} يوماً.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("المادة (7): القانون المطبق")
                                    doc.add_paragraph("تخضع هذه الاتفاقية لأنظمة المملكة العربية السعودية.")
                                    doc.add_paragraph()
                                    doc.add_paragraph("التوقيع:")
                                    doc.add_paragraph("مقدم الخدمة: _____________________")
                                    doc.add_paragraph("العميل: _________________________")
                            elif template.category == 'Commercial':
                                # Full Commercial Sale Contract template
                                doc.add_paragraph("عقد بيع تجاري")
                                doc.add_paragraph()
                                doc.add_paragraph(f"تم بعون الله بتاريخ {{{{ purchase_date }}}} بين كل من:")
                                doc.add_paragraph()
                                doc.add_paragraph(f"1. {{{{ seller_name }}}}، ويشار إليه بـ \"البائع\".")
                                doc.add_paragraph(f"2. {{{{ buyer_name }}}}، ويشار إليه بـ \"المشتري\".")
                                doc.add_paragraph()
                                doc.add_paragraph("المادة (1): موضوع العقد")
                                doc.add_paragraph(f"باع البائع للمشتري البضاعة التالية: {{{{ item_description }}}}.")
                                doc.add_paragraph()
                                doc.add_paragraph("المادة (2): السعر وطريقة الدفع")
                                doc.add_paragraph(f"تم الاتفاق على أن يكون سعر البضاعة {{{{ purchase_amount }}}} ريال سعودي، يتم سداده نقداً أو عن طريق التحويل البنكي خلال {{{{ payment_terms }}}} يوم عمل.")
                                doc.add_paragraph()
                                doc.add_paragraph("المادة (3): التسليم")
                                doc.add_paragraph(f"يتم تسليم البضاعة في موقع {{{{ delivery_location }}}} بتاريخ {{{{ delivery_date }}}}.")
                                doc.add_paragraph()
                                doc.add_paragraph("المادة (4): الضمان")
                                doc.add_paragraph(f"يضمن البائع خلو البضاعة من العيوب لمدة {{{{ warranty_period }}}} أيام من تاريخ التسليم.")
                                doc.add_paragraph()
                                doc.add_paragraph("المادة (5): القانون المطبق")
                                doc.add_paragraph("يخضع هذا العقد لأنظمة المملكة العربية السعودية.")
                                doc.add_paragraph()
                                doc.add_paragraph("التوقيع:")
                                doc.add_paragraph("البائع: _____________________")
                                doc.add_paragraph("المشتري: ___________________")
                            else:
                                # Generic template
                                for var in variables_data:
                                    if isinstance(var, dict) and 'name' in var and 'label' in var:
                                        label = arabic_labels.get(var['label'], var['label'])
                                        doc.add_paragraph(f"{label}: {{{{ {var['name']} }}}}")
                            
                            doc.add_paragraph()
                            doc.add_paragraph("يتفق الطرفان على الشروط والأحكام الواردة أعلاه.")
                            doc.add_paragraph()
                            doc.add_paragraph("التوقيع: _____________________    التاريخ: _____________________")
                            
                            # Save document
                            doc.save(str(new_file_path))
                            
                            # Update template file path
                            template.file_path = str(new_file_path)
                            template.format = 'docx'
                            
                            await db.flush()
                            
                            print(f"  ✅ Created DOCX file: {new_file_path}")
                            
                        except ImportError:
                            print(f"  ⚠️  python-docx not available, skipping template: {template.title}")
                            continue
                        except Exception as e:
                            print(f"  ❌ Error creating DOCX for {template.title}: {str(e)}")
                            continue
                    else:
                        print(f"  ✓ Template already has valid file: {template.title}")
                    
                except Exception as e:
                    print(f"  ❌ Error processing template {template.title}: {str(e)}")
                    continue
            
            await db.commit()
            
            print("\n" + "=" * 60)
            print("✨ Template files fixed successfully!")
            print("=" * 60)
            
        except Exception as e:
            print(f"\n❌ Error: {str(e)}")
            import traceback
            traceback.print_exc()
            await db.rollback()
            raise


if __name__ == "__main__":
    asyncio.run(fix_template_files())

